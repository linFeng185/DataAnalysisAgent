"""数据源管理、Schema 刷新、上传限制与 XSS 回归测试。"""

from __future__ import annotations

import io
import sys
from types import SimpleNamespace
from unittest.mock import AsyncMock

import pytest
from fastapi import HTTPException, UploadFile


class TestExtensionScopeAuthorization:
    """覆盖功能 8.4.1、9.5.6：Skill/MCP 三级作用域管理授权。"""

    # 方法作用：验证租户管理员不能把上传 Skill 写入系统作用域。
    # Args: self - pytest 测试类实例；monkeypatch - pytest 补丁工具；tmp_path - 临时目录。
    # Returns: 无返回值，断言失败时由 pytest 报告。
    async def test_tenant_admin_cannot_upload_system_skill(self, monkeypatch, tmp_path):
        """系统 Skill 只能由平台超级管理员或配置目录扫描任务维护。"""
        # Arrange
        import src.api.auth as auth
        import src.api.routes as routes
        import src.skill_manager as skill_module

        manager = skill_module.SkillManager(
            str(tmp_path / "builtin"), managed_dir=str(tmp_path / "managed"),
        )
        monkeypatch.setattr(skill_module, "get_skill_manager", lambda *args, **kwargs: manager)
        monkeypatch.setattr(auth, "get_current_role", lambda: "tenant_admin")
        monkeypatch.setattr(auth, "get_current_tenant_id", lambda: 4)
        monkeypatch.setattr(auth, "get_current_user_id", lambda: 7)
        upload = UploadFile(
            filename="report/SKILL.md",
            file=io.BytesIO(b"---\nname: report\nversion: 1.0.0\n---\nbody"),
        )

        # Act / Assert
        with pytest.raises(HTTPException) as caught:
            await routes.upload_skills([upload], skill_scope="system")
        assert caught.value.status_code == 403
        assert not (tmp_path / "builtin").exists()

    # 方法作用：验证普通用户不能创建租户公共 MCP 服务。
    # Args: self - pytest 测试类实例；monkeypatch - pytest 补丁工具。
    # Returns: 无返回值，断言失败时由 pytest 报告。
    async def test_analyst_cannot_create_tenant_mcp_server(self, monkeypatch):
        """tenant MCP 的写权限仅属于租户管理员或平台超级管理员。"""
        # Arrange
        import src.api.auth as auth
        import src.api.routes as routes
        from src.api.schemas import MCPServerCreate

        monkeypatch.setattr(auth, "get_current_role", lambda: "analyst")
        monkeypatch.setattr(auth, "get_current_tenant_id", lambda: 4)
        monkeypatch.setattr(auth, "get_current_user_id", lambda: 7)
        request = MCPServerCreate(
            name="tenant-docs", transport="sse", url="http://127.0.0.1:9000/sse",
            scope="tenant",
        )

        # Act / Assert
        with pytest.raises(HTTPException) as caught:
            await routes.create_mcp_server(request)
        assert caught.value.status_code == 403

    # 方法作用：验证 MCP 请求模型保留显式作用域。
    # Args: self - pytest 测试类实例。
    # Returns: 无返回值，断言失败时由 pytest 报告。
    def test_mcp_create_schema_defaults_to_private_scope(self):
        """未指定作用域的 MCP 配置默认仅当前用户可见。"""
        # Arrange / Act
        from src.api.schemas import MCPServerCreate

        request = MCPServerCreate(name="personal-files")

        # Assert
        assert request.scope == "private"

    # 方法作用：验证普通用户上传 Skill 默认进入本人私有目录并注入可见缓存。
    # Args: self - pytest 测试类实例；monkeypatch - pytest 补丁工具；tmp_path - 临时目录。
    # Returns: 无返回值，断言失败时由 pytest 报告。
    async def test_private_skill_upload_uses_current_identity(self, monkeypatch, tmp_path):
        """上传包不能伪造作用域，最终路径由当前 tenant_id/user_id 决定。"""
        # Arrange
        import src.api.auth as auth
        import src.api.routes as routes
        import src.config as config_module
        import src.skill_manager as skill_module

        manager = skill_module.SkillManager(
            str(tmp_path / "builtin"), managed_dir=str(tmp_path / "managed"),
        )
        monkeypatch.setattr(skill_module, "get_skill_manager", lambda *args, **kwargs: manager)
        monkeypatch.setattr(auth, "get_current_role", lambda: "analyst")
        monkeypatch.setattr(auth, "get_current_tenant_id", lambda: 4)
        monkeypatch.setattr(auth, "get_current_user_id", lambda: 7)
        monkeypatch.setattr(
            config_module, "get_settings", lambda: SimpleNamespace(multi_tenant=True),
        )
        upload = UploadFile(
            filename="report/SKILL.md",
            file=io.BytesIO(
                b"---\nname: report\nversion: 1.0.0\ntriggers:\n  keywords: [report]\n---\nbody",
            ),
        )

        # Act
        result = await routes.upload_skills([upload], skill_scope="private")

        # Assert
        expected = tmp_path / "managed" / "private" / "4" / "7" / "report" / "SKILL.md"
        assert result["total"] == 1
        assert expected.is_file()
        skill = manager.get_skill("report", scope="private", tenant_id=4, user_id=7)
        assert skill is not None
        assert skill.owner_user_id == 7

    # 方法作用：验证个人 MCP 创建只写入当前租户和当前用户参数。
    # Args: self - pytest 测试类实例；monkeypatch - pytest 补丁工具。
    # Returns: 无返回值，断言失败时由 pytest 报告。
    async def test_private_mcp_create_uses_current_identity(self, monkeypatch):
        """API 不接受请求体中的所有者，数据库参数必须来自认证上下文。"""
        # Arrange
        import src.api.auth as auth
        import src.api.routes as routes
        import src.config as config_module
        import src.mcp_client.client_manager as mcp_module
        from src.api.schemas import MCPServerCreate

        connection = SimpleNamespace(
            fetchval=AsyncMock(return_value=None), execute=AsyncMock(), close=AsyncMock(),
        )
        monkeypatch.setitem(
            sys.modules, "asyncpg",
            SimpleNamespace(connect=AsyncMock(return_value=connection)),
        )
        manager = SimpleNamespace(ensure_scoped_servers=AsyncMock(return_value=1))
        monkeypatch.setattr(mcp_module, "get_mcp_client_manager", lambda: manager)
        monkeypatch.setattr(auth, "get_current_role", lambda: "analyst")
        monkeypatch.setattr(auth, "get_current_tenant_id", lambda: 4)
        monkeypatch.setattr(auth, "get_current_user_id", lambda: 7)
        monkeypatch.setattr(
            config_module, "get_settings",
            lambda: SimpleNamespace(
                multi_tenant=True, database_url="postgresql+asyncpg://test:test@db/test",
            ),
        )
        request = MCPServerCreate(
            name="personal-files", scope="private", transport="sse",
            url="http://127.0.0.1:9000/sse",
        )

        # Act
        result = await routes.create_mcp_server(request)

        # Assert
        assert result["scope"] == "private"
        insert_args = connection.execute.await_args.args
        assert insert_args[3:5] == (4, 7)
        manager.ensure_scoped_servers.assert_awaited_once_with(4, 7, force=True)


class TestDatasourceLifecycle:
    """覆盖注册、列表、删除的真实 Registry 生命周期。"""

    async def test_register_is_visible_and_delete_removes_provider(self, monkeypatch):
        """注册后的数据源应进入全局列表，删除后不可见。"""
        # Arrange
        import src.api.routes as routes
        from src.datasource.providers.external import ExternalDataSourceProvider
        from src.datasource.registry import DataSourceRegistry

        provider = ExternalDataSourceProvider()
        monkeypatch.setattr(provider, "_prefetch_schema", AsyncMock())
        registry = DataSourceRegistry()
        registry.register_provider("external", provider)
        monkeypatch.setattr(routes, "_registry", lambda: registry)
        request = routes.DataSourceCreateRequest(
            name="managed", dialect="sqlite", file_path=":memory:",
        )

        # Act
        created = await routes.register_datasource(request)
        listed_after_create = await routes.list_datasources(page=1, page_size=100)
        deleted = await routes.delete_datasource("managed")
        listed_after_delete = await routes.list_datasources(page=1, page_size=100)

        # Assert
        assert created.name == "managed"
        assert any(item["name"] == "managed" for item in listed_after_create["datasources"])
        assert deleted["status"] == "ok"
        assert not any(item["name"] == "managed" for item in listed_after_delete["datasources"])

    async def test_delete_unknown_datasource_returns_404(self, monkeypatch):
        """删除不存在的数据源必须返回 404 而非假成功。"""
        # Arrange
        import src.api.routes as routes
        from src.datasource.registry import DataSourceRegistry

        registry = DataSourceRegistry()
        monkeypatch.setattr(routes, "_registry", lambda: registry)

        # Act / Assert
        with pytest.raises(HTTPException) as caught:
            await routes.delete_datasource("missing")
        assert caught.value.status_code == 404

    async def test_asgi_datasource_lifecycle_changes_status(self, monkeypatch):
        """ASGI 请求应真实反映注册和删除后的 Registry 状态。"""
        # Arrange
        from httpx import ASGITransport, AsyncClient

        import src.api.routes as routes
        from src.datasource.providers.external import ExternalDataSourceProvider
        from src.datasource.registry import DataSourceRegistry
        from src.main import create_app

        registry = DataSourceRegistry()
        registry.register_provider("external", ExternalDataSourceProvider())
        monkeypatch.setattr(routes, "_registry", lambda: registry)
        client = AsyncClient(transport=ASGITransport(app=create_app()), base_url="http://test")

        # Act
        created = await client.post("/api/v1/datasources", json={
            "name": "http-managed",
            "dialect": "sqlite",
            "file_path": ":memory:",
        })
        listed = await client.get("/api/v1/datasources")
        deleted = await client.delete("/api/v1/datasources/http-managed")
        deleted_again = await client.delete("/api/v1/datasources/http-managed")
        await client.aclose()

        # Assert
        assert created.status_code == 201
        assert any(item["name"] == "http-managed" for item in listed.json()["datasources"])
        assert deleted.status_code == 200
        assert deleted_again.status_code == 404


class TestSchemaManagement:
    """覆盖 Schema 刷新和字段备注写回。"""

    async def test_refresh_unknown_datasource_returns_404(self, monkeypatch):
        """Schema 刷新必须验证数据源真实存在。"""
        # Arrange
        import src.api.routes as routes
        from src.datasource.registry import DataSourceRegistry

        registry = DataSourceRegistry()
        monkeypatch.setattr(routes, "_registry", lambda: registry)

        # Act / Assert
        with pytest.raises(HTTPException) as caught:
            await routes.refresh_schema(datasource="missing")
        assert caught.value.status_code == 404

    async def test_comment_delegates_to_schema_manager(self, monkeypatch):
        """字段备注更新必须写入 SchemaManager，而不是只返回成功消息。"""
        # Arrange
        import src.api.routes as routes

        manager = SimpleNamespace(update_column_comment=AsyncMock(return_value=True))
        monkeypatch.setattr(routes, "_registry", lambda: SimpleNamespace(resolve=AsyncMock(return_value=object())))
        monkeypatch.setattr(routes, "_schema_manager", lambda: manager, raising=False)
        request = routes.ColumnCommentRequest(comment="订单金额")

        # Act
        result = await routes.update_column_comment("orders", "amount", request, datasource="demo")

        # Assert
        assert result["status"] == "ok"
        manager.update_column_comment.assert_awaited_once_with("demo", "orders", "amount", "订单金额")


class TestKnowledgeUploadSafety:
    """覆盖文档 XSS 和上传大小限制。"""

    def test_docx_html_escapes_text_and_cells(self):
        """Word 段落和表格单元格的 HTML 特殊字符必须转义。"""
        # Arrange
        from docx import Document
        import src.api.routes as routes

        document = Document()
        document.add_paragraph("<script>alert(1)</script>")
        table = document.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "<img src=x onerror=alert(1)>"
        buffer = io.BytesIO()
        document.save(buffer)

        # Act
        html = routes._docx_to_html(buffer.getvalue())

        # Assert
        assert "<script>" not in html
        assert "<img" not in html
        assert "&lt;script&gt;" in html
        assert "&lt;img" in html

    async def test_upload_rejects_content_over_limit(self, monkeypatch):
        """上传内容超过 MAX_UPLOAD_BYTES 时应在处理前返回 413。"""
        # Arrange
        import src.config as config_module
        import src.api.routes as routes

        monkeypatch.setattr(config_module, "get_settings", lambda: SimpleNamespace(max_upload_bytes=4))
        upload = UploadFile(filename="too-large.txt", file=io.BytesIO(b"12345"))

        # Act / Assert
        with pytest.raises(HTTPException) as caught:
            await routes.upload_knowledge_docs(
                [upload], strategy="auto", chunk_size=800, chunk_overlap=100, category="",
            )
        assert caught.value.status_code == 413


class TestKnowledgeVectorStoreRoutes:
    """覆盖 6.8.6 知识管理 API 的 VectorStore 和租户边界。"""

    async def test_list_knowledge_uses_scoped_vector_store_filter(self, monkeypatch):
        """知识列表应分别读取系统、租户和个人范围，而不是执行宽查询。"""
        # Arrange
        import src.api.routes as routes
        import src.memory.vector_store as vector_module
        from src.memory.vector_store import VectorEntry

        store = SimpleNamespace(get_by_filter=AsyncMock(side_effect=[
            [],
            [VectorEntry(
                id="doc-1", content="GMV 定义",
                metadata={"category": "metric", "source": "user_upload", "visibility": "tenant"},
            )],
            [],
        ]))
        monkeypatch.setattr(vector_module, "get_vector_store", AsyncMock(return_value=store))

        # Act
        result = await routes.list_knowledge(category="metric", search=None, page=1, page_size=20)

        # Assert
        assert result["total"] == 1
        assert result["entries"][0]["id"] == "doc-1"
        filters = [call.args[0] for call in store.get_by_filter.await_args_list]
        assert [item["visibility"] for item in filters] == ["system", "tenant", "private"]
        assert all(item["category"] == "metric" for item in filters)

    async def test_delete_knowledge_entry_checks_owner_before_delete(self, monkeypatch):
        """删除用户知识条目前必须校验租户和所有者。"""
        # Arrange
        import src.api.routes as routes
        import src.memory.vector_store as vector_module
        import src.config as config_module
        from src.memory.vector_store import VectorEntry

        store = SimpleNamespace(
            get_by_id=AsyncMock(return_value=VectorEntry(
                id="doc-2", content="私有知识",
                metadata={"source": "user_upload", "tenant_id": 4, "owner_user_id": 7},
            )),
            delete_by_ids=AsyncMock(return_value=1),
        )
        monkeypatch.setattr(vector_module, "get_vector_store", AsyncMock(return_value=store))
        monkeypatch.setattr(config_module, "get_settings", lambda: SimpleNamespace(multi_tenant=True))
        monkeypatch.setattr("src.api.auth.get_current_tenant_id", lambda: 4)
        monkeypatch.setattr("src.api.auth.get_current_user_id", lambda: 7)

        # Act
        result = await routes.delete_knowledge_entry("doc-2")

        # Assert
        assert result == {"status": "ok", "id": "doc-2"}
        store.delete_by_ids.assert_awaited_once_with(["doc-2"])

    async def test_tenant_admin_can_delete_tenant_shared_entry(self, monkeypatch):
        """租户管理员可治理本租户公共知识，不受上传者所有权限制。"""
        # Arrange
        import src.api.routes as routes
        import src.memory.vector_store as vector_module
        import src.config as config_module
        from src.memory.vector_store import VectorEntry

        store = SimpleNamespace(
            get_by_id=AsyncMock(return_value=VectorEntry(
                id="tenant-doc", content="租户口径",
                metadata={
                    "source": "user_upload", "visibility": "tenant",
                    "tenant_id": 4, "uploaded_by_user_id": 99,
                },
            )),
            delete_by_ids=AsyncMock(return_value=1),
        )
        monkeypatch.setattr(vector_module, "get_vector_store", AsyncMock(return_value=store))
        monkeypatch.setattr(config_module, "get_settings", lambda: SimpleNamespace(multi_tenant=True))
        monkeypatch.setattr("src.api.auth.get_current_tenant_id", lambda: 4)
        monkeypatch.setattr("src.api.auth.get_current_user_id", lambda: 7)
        monkeypatch.setattr("src.api.auth.get_current_role", lambda: "tenant_admin")

        # Act
        result = await routes.delete_knowledge_entry("tenant-doc")

        # Assert
        assert result == {"status": "ok", "id": "tenant-doc"}
        store.delete_by_ids.assert_awaited_once_with(["tenant-doc"])

    async def test_tenant_admin_can_delete_tenant_shared_document(self, monkeypatch):
        """租户管理员删除公共文档时应同时清理原文件和对应向量。"""
        # Arrange
        import src.api.routes as routes
        import src.memory.vector_store as vector_module
        import src.knowledge.file_store as file_module
        from src.memory.vector_store import VectorEntry

        entry = VectorEntry(
            id="tenant:4:chunk", content="租户口径",
            metadata={"source": "user_upload", "visibility": "tenant", "tenant_id": 4},
        )
        store = SimpleNamespace(
            get_by_filter=AsyncMock(side_effect=[[], [entry], []]),
            delete_by_ids=AsyncMock(return_value=1),
        )
        file_store = SimpleNamespace(delete=AsyncMock(return_value=True))
        monkeypatch.setattr(vector_module, "get_vector_store", AsyncMock(return_value=store))
        monkeypatch.setattr(file_module, "get_file_store", lambda: file_store)
        monkeypatch.setattr("src.api.auth.get_current_tenant_id", lambda: 4)
        monkeypatch.setattr("src.api.auth.get_current_user_id", lambda: 7)
        monkeypatch.setattr("src.api.auth.get_current_role", lambda: "tenant_admin")

        # Act
        result = await routes.delete_knowledge_doc("tenant.md", knowledge_scope="tenant")

        # Assert
        assert result == {"status": "ok", "doc": "tenant.md"}
        file_store.delete.assert_awaited_once_with("tenant.md", knowledge_scope="tenant")
        store.delete_by_ids.assert_awaited_once_with(["tenant:4:chunk"])


class TestDataIntelligenceRoutes:
    """覆盖 Phase B/D 结构化 profile 和预测 API。"""

    async def test_profile_structured_asset_csv(self, monkeypatch):
        """结构化资产 profile API 应返回列级质量信息。"""
        # Arrange
        import src.api.routes as routes
        import src.config as config_module

        monkeypatch.setattr(config_module, "get_settings", lambda: SimpleNamespace(max_upload_bytes=10000))
        upload = UploadFile(filename="orders.csv", file=io.BytesIO(b"id,value\n1,2\n2,3\n"))

        # Act
        result = await routes.profile_structured_asset(upload)

        # Assert
        assert result["format"] == "csv"
        assert result["row_count"] == 2
        assert result["columns"]["id"]["unique_count"] == 2

    async def test_forecast_asset_returns_model_card(self):
        """预测 API 应返回模型卡和回测指标，而不是裸预测数组。"""
        # Arrange
        import src.api.routes as routes
        rows = [{"date": f"2026-01-{index:02d}", "value": index} for index in range(1, 9)]

        # Act
        result = await routes.forecast_asset({
            "rows": rows, "time_col": "date", "value_col": "value", "horizon": 2,
        })

        # Assert
        assert len(result["predictions"]) == 2
        assert result["model_card"]["leakage_check"] == "passed"

    async def test_query_structured_asset_requires_duckdb(self, monkeypatch):
        """结构化 SQL API 缺少 DuckDB 时应返回明确的 400。"""
        # Arrange
        import src.api.routes as routes
        import src.config as config_module
        import src.knowledge.structured_query as query_module

        monkeypatch.setattr(config_module, "get_settings", lambda: SimpleNamespace(
            max_upload_bytes=10000, max_result_rows=100,
        ))
        monkeypatch.setattr(query_module, "_load_duckdb", lambda: None)
        upload = UploadFile(filename="orders.csv", file=io.BytesIO(b"id\n1\n"))

        # Act / Assert
        with pytest.raises(HTTPException) as caught:
            await routes.query_structured_asset(upload, sql="SELECT * FROM data", sheet_name=None)
        assert caught.value.status_code == 400
        assert "DuckDB" in str(caught.value.detail)
