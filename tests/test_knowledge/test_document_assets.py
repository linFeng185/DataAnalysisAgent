"""Phase C 文档资产结构与引用定位测试。"""

from __future__ import annotations

import io

import pytest


class TestDocumentAssetAdapter:
    """覆盖 Markdown/Word 文档的结构保真解析。"""

    def test_parse_markdown_keeps_paragraph_locators(self):
        """Markdown 应按段落生成可引用的行定位。"""
        # Arrange
        from src.knowledge.document_assets import DocumentAssetAdapter

        content = "# 指标\n\nGMV 是成交总额。\n\n## 口径\n\n不含退款。".encode()

        # Act
        asset = DocumentAssetAdapter().parse("metric.md", content)

        # Assert
        assert asset.mime_type == "text/markdown"
        assert len(asset.blocks) == 4
        assert asset.blocks[0].locator["line_start"] == 1
        assert asset.blocks[1].heading_path == ["指标"]
        assert asset.blocks[3].heading_path == ["指标", "口径"]

    def test_parse_docx_keeps_paragraph_and_table_locators(self):
        """Word 应保留段落样式和表格行定位，而不是只输出纯文本。"""
        # Arrange
        from docx import Document
        from src.knowledge.document_assets import DocumentAssetAdapter

        doc = Document()
        doc.add_heading("销售报告", level=1)
        doc.add_paragraph("本月销售额 100 万。")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "渠道"
        table.cell(0, 1).text = "金额"
        table.cell(1, 0).text = "线上"
        table.cell(1, 1).text = "60"
        buffer = io.BytesIO()
        doc.save(buffer)

        # Act
        asset = DocumentAssetAdapter().parse("report.docx", buffer.getvalue())

        # Assert
        assert asset.mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        assert any(block.locator.get("paragraph") == 1 for block in asset.blocks)
        assert any(block.locator.get("table") == 0 and block.locator.get("row") == 1 for block in asset.blocks)
        assert asset.blocks[0].heading_path == ["销售报告"]

    def test_chunk_document_propagates_citation_locator(self):
        """分块后必须保留原文 locator，供 Evidence/Citation 回溯。"""
        # Arrange
        from src.knowledge.document_assets import DocumentAssetAdapter, chunk_document
        from src.knowledge.doc_parser import ChunkConfig, ChunkStrategy

        asset = DocumentAssetAdapter().parse("note.txt", "第一段内容\n\n第二段内容".encode("utf-8"))

        # Act
        chunks = chunk_document(asset, ChunkConfig(
            strategy=ChunkStrategy.PARAGRAPH, min_chunk_size=1, chunk_size=100,
        ))

        # Assert
        assert chunks
        assert chunks[0].metadata["locator"]["line_start"] == 1

    def test_parse_unsupported_format(self):
        """不支持的文档格式必须拒绝。"""
        from src.knowledge.document_assets import DocumentAssetAdapter, DocumentAssetError

        with pytest.raises(DocumentAssetError, match="不支持"):
            DocumentAssetAdapter().parse("data.exe", b"x")
