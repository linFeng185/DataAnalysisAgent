"""结构保真文档解析与证据定位。"""

from __future__ import annotations

import hashlib
import html
import io
import os
import re
from dataclasses import dataclass, field
from typing import Any

from src.knowledge.doc_parser import ChunkConfig, ChunkStrategy, DocChunk
from src.logging_config import get_logger

logger = get_logger(__name__)


class DocumentAssetError(ValueError):
    """文档资产解析失败时抛出的异常。"""


@dataclass
class DocumentBlock:
    """带原文定位的文档块。"""

    content: str
    locator: dict[str, Any]
    heading_path: list[str] = field(default_factory=list)


@dataclass
class DocumentAsset:
    """文档原件的可检索结构和版本信息。"""

    asset_id: str
    source_file: str
    mime_type: str
    checksum: str
    blocks: list[DocumentBlock]
    page_count: int = 0
    warnings: list[str] = field(default_factory=list)


class DocumentAssetAdapter:
    """解析 Markdown/TXT/HTML/PDF/Word 并保留段落、页码或表格定位。"""

    _MIME = {
        ".txt": "text/plain",
        ".md": "text/markdown",
        ".markdown": "text/markdown",
        ".html": "text/html",
        ".htm": "text/html",
        ".pdf": "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }

    # 方法作用：初始化文档适配器并限制单个文档的读取大小。
    # Args: max_bytes - 允许解析的最大字节数。
    # Returns: 无返回值。
    def __init__(self, max_bytes: int = 100 * 1024 * 1024) -> None:
        logger.debug("文档资产适配器初始化入口", max_bytes=max_bytes)
        if max_bytes <= 0:
            raise ValueError("文档大小上限必须大于零")
        self.max_bytes = max_bytes
        logger.info("文档资产适配器初始化完成", max_bytes=max_bytes)

    # 方法作用：按文件扩展名解析文档并生成带 checksum 的统一资产。
    # Args: file_name - 原始文件名；content - 文档二进制内容。
    # Returns: DocumentAsset 文档资产。
    def parse(self, file_name: str, content: bytes) -> DocumentAsset:
        logger.debug("文档资产解析入口", file_name=file_name, content_size=len(content))
        if len(content) > self.max_bytes:
            raise DocumentAssetError(f"文档大小超过限制 {self.max_bytes} 字节")
        if not content:
            raise DocumentAssetError("文档内容为空")
        ext = os.path.splitext(file_name)[1].lower()
        mime_type = self._MIME.get(ext)
        if not mime_type:
            raise DocumentAssetError(f"不支持的文档格式: {ext or file_name}")
        try:
            if ext in {".txt", ".md", ".markdown"}:
                blocks = self._parse_text(content.decode("utf-8", errors="replace"), markdown=ext != ".txt")
                page_count = 0
            elif ext in {".html", ".htm"}:
                text = self._strip_html(content.decode("utf-8", errors="replace"))
                blocks = self._parse_text(text, markdown=False)
                page_count = 0
            elif ext == ".pdf":
                blocks, page_count = self._parse_pdf(content)
            else:
                blocks = self._parse_docx(content)
                page_count = 0
            if not blocks:
                raise DocumentAssetError("文档未提取到正文")
            checksum = hashlib.sha256(content).hexdigest()
            asset = DocumentAsset(
                asset_id=f"document:{checksum[:24]}",
                source_file=file_name,
                mime_type=mime_type,
                checksum=checksum,
                blocks=blocks,
                page_count=page_count,
            )
            logger.info("文档资产解析完成", file_name=file_name, blocks=len(blocks), pages=page_count)
            return asset
        except DocumentAssetError:
            logger.error("文档资产解析失败", file_name=file_name, exc_info=True)
            raise
        except ImportError as exc:
            logger.error("文档可选解析引擎缺失", file_name=file_name, error=str(exc), exc_info=True)
            raise DocumentAssetError(f"{ext} 解析需要可选依赖: {exc}") from exc
        except Exception as exc:
            logger.error("文档资产解析异常", file_name=file_name, error=str(exc), exc_info=True)
            raise DocumentAssetError(f"{file_name} 解析失败: {exc}") from exc

    # 方法作用：解析文本或 Markdown，保留段落行号和标题层级。
    # Args: text - 解码后的文本；markdown - 是否识别 Markdown 标题。
    # Returns: 带定位的 DocumentBlock 列表。
    def _parse_text(self, text: str, markdown: bool) -> list[DocumentBlock]:
        logger.debug("解析文本块入口", text_size=len(text), markdown=markdown)
        lines = text.splitlines()
        blocks: list[DocumentBlock] = []
        headings: list[str] = []
        paragraph: list[str] = []
        paragraph_start = 1

        # 方法作用：把累计段落写成一个带行范围的块。
        # Args: 无；使用外层 paragraph、paragraph_start 和 headings。
        # Returns: 无返回值。
        def flush_paragraph() -> None:
            nonlocal paragraph
            logger.debug("刷新文本段落入口", pending_lines=len(paragraph))
            if paragraph and "\n".join(paragraph).strip():
                blocks.append(DocumentBlock(
                    content="\n".join(paragraph).strip(),
                    locator={"line_start": paragraph_start, "line_end": paragraph_start + len(paragraph) - 1},
                    heading_path=list(headings),
                ))
            paragraph = []
            logger.info("刷新文本段落完成", block_count=len(blocks))

        for index, line in enumerate(lines, start=1):
            heading_match = re.match(r"^(#{1,6})\s+(.+?)\s*$", line) if markdown else None
            if heading_match:
                flush_paragraph()
                level = len(heading_match.group(1))
                title = heading_match.group(2)
                headings[:] = headings[:level - 1]
                headings.append(title)
                blocks.append(DocumentBlock(
                    content=title,
                    locator={"line_start": index, "line_end": index},
                    heading_path=list(headings),
                ))
                paragraph_start = index + 1
            elif line.strip():
                if not paragraph:
                    paragraph_start = index
                paragraph.append(line)
            else:
                flush_paragraph()
                paragraph_start = index + 1
        flush_paragraph()
        logger.info("解析文本块完成", block_count=len(blocks))
        return blocks

    # 方法作用：解析 PDF 每页文字并把页码写入 locator。
    # Args: content - PDF 二进制内容。
    # Returns: (DocumentBlock 列表, 页数)。
    def _parse_pdf(self, content: bytes) -> tuple[list[DocumentBlock], int]:
        logger.debug("解析 PDF 入口", content_size=len(content))
        from PyPDF2 import PdfReader

        reader = PdfReader(io.BytesIO(content))
        blocks: list[DocumentBlock] = []
        for page_number, page in enumerate(reader.pages, start=1):
            text = (page.extract_text() or "").strip()
            if text:
                blocks.append(DocumentBlock(content=text, locator={"page": page_number}))
        logger.info("解析 PDF 完成", pages=len(reader.pages), blocks=len(blocks))
        return blocks, len(reader.pages)

    # 方法作用：解析 Word 段落和表格，保留样式、表号和行定位。
    # Args: content - DOCX 二进制内容。
    # Returns: 带结构定位的 DocumentBlock 列表。
    def _parse_docx(self, content: bytes) -> list[DocumentBlock]:
        logger.debug("解析 Word 入口", content_size=len(content))
        from docx import Document

        doc = Document(io.BytesIO(content))
        blocks: list[DocumentBlock] = []
        headings: list[str] = []
        for index, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            style_name = para.style.name if para.style else ""
            if style_name.startswith("Heading"):
                match = re.search(r"(\d+)", style_name)
                level = int(match.group(1)) if match else 1
                headings[:] = headings[:level - 1]
                headings.append(text)
            blocks.append(DocumentBlock(
                content=text,
                locator={"paragraph": index, "style": style_name or "Normal"},
                heading_path=list(headings),
            ))
        for table_index, table in enumerate(doc.tables):
            for row_index, row in enumerate(table.rows):
                values = [cell.text.strip() for cell in row.cells]
                if not any(values):
                    continue
                blocks.append(DocumentBlock(
                    content=" | ".join(values),
                    locator={"table": table_index, "row": row_index,
                             "cell_range": f"R{row_index + 1}:C1-C{len(values)}"},
                    heading_path=list(headings),
                ))
        logger.info("解析 Word 完成", paragraphs=len(doc.paragraphs), tables=len(doc.tables), blocks=len(blocks))
        return blocks

    # 方法作用：去除 HTML 标签并解码实体，形成可检索正文。
    # Args: text - HTML 文本。
    # Returns: 纯文本内容。
    @staticmethod
    def _strip_html(text: str) -> str:
        logger.debug("清理 HTML 入口", text_size=len(text))
        without_script = re.sub(r"<\s*(script|style)[^>]*>.*?<\s*/\s*\1\s*>", "", text, flags=re.I | re.S)
        plain = re.sub(r"<[^>]+>", "\n", without_script)
        result = html.unescape(re.sub(r"\n{3,}", "\n\n", plain)).strip()
        logger.info("清理 HTML 完成", output_size=len(result))
        return result


# 方法作用：将结构化文档块按分块策略转换为可写入 VectorStore 的 DocChunk。
# Args: asset - DocumentAsset；config - 分块大小、重叠和策略配置。
# Returns: 带 locator 元数据的 DocChunk 列表。
def chunk_document(asset: DocumentAsset, config: ChunkConfig) -> list[DocChunk]:
    logger.debug("文档资产分块入口", file_name=asset.source_file, blocks=len(asset.blocks))
    strategy = config.strategy
    if strategy == ChunkStrategy.AUTO:
        strategy = ChunkStrategy.HEADING if any(block.heading_path for block in asset.blocks) else ChunkStrategy.PARAGRAPH
    chunks: list[DocChunk] = []
    current: list[DocumentBlock] = []
    current_size = 0

    # 方法作用：把累计文档块合并成一个可引用分块。
    # Args: blocks - 要合并的文档块；chunk_index - 分块序号。
    # Returns: 生成的 DocChunk，内容为空时返回 None。
    def make_chunk(blocks: list[DocumentBlock], chunk_index: int) -> DocChunk | None:
        logger.debug("生成文档分块入口", chunk_index=chunk_index, block_count=len(blocks))
        if not blocks:
            return None
        content = "\n\n".join(block.content for block in blocks).strip()
        if len(content) < config.min_chunk_size:
            logger.info("文档分块因最小长度跳过", chunk_index=chunk_index)
            return None
        first, last = blocks[0], blocks[-1]
        locator = dict(first.locator)
        if last.locator != first.locator:
            locator["end"] = dict(last.locator)
        result = DocChunk(
            content=content,
            metadata={
                "source_file": asset.source_file,
                "strategy": strategy.value,
                "chunk_size": len(content),
                "locator": locator,
                "heading_path": list(first.heading_path),
                "checksum": asset.checksum,
                "document_version": "v1",
            },
        )
        logger.info("生成文档分块完成", chunk_index=chunk_index, content_size=len(content))
        return result

    for block in asset.blocks:
        if strategy == ChunkStrategy.HEADING and current and block.heading_path != current[-1].heading_path:
            chunk = make_chunk(current, len(chunks))
            if chunk:
                chunks.append(chunk)
            current, current_size = [], 0
        if current and current_size + len(block.content) + 2 > config.chunk_size:
            chunk = make_chunk(current, len(chunks))
            if chunk:
                chunks.append(chunk)
            current, current_size = [], 0
        current.append(block)
        current_size += len(block.content) + 2
    chunk = make_chunk(current, len(chunks))
    if chunk:
        chunks.append(chunk)
    logger.info("文档资产分块完成", file_name=asset.source_file, chunks=len(chunks), strategy=strategy.value)
    return chunks
