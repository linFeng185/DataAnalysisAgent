"""知识文档、标签与检索管理路由。"""

from __future__ import annotations

import io
import html
import json
import os
import time
import uuid
from collections.abc import AsyncIterator
from contextlib import asynccontextmanager

from fastapi import APIRouter, Body, File, HTTPException, Query, UploadFile

from src.api.schemas import (
    ChatRequest, ChatResponse, ColumnCommentRequest,
    DataSourceCreateRequest, DataSourceInfo, HealthResponse, KnowledgeTagCreateRequest,
    KnowledgeTagStatusRequest, MCPServerCreate, TableInfo,
)
from src.exceptions import DataSourceNotFoundError
from src.llm.client import is_llm_available
from src.logging_config import get_logger
from src.api.routes._helpers import _app, _authorize_extension_scope, _registry

logger = get_logger(__name__)
router = APIRouter()
_started_at = time.time()



def _knowledge_where(extra: dict | None = None, owner_only: bool = False) -> dict | None:
    """构建兼容 ChromaDB 的知识库租户过滤条件。

    Args:
        extra: 额外 metadata 精确过滤条件。
        owner_only: 是否同时限制为当前用户创建。

    Returns:
        ChromaDB where 条件；单租户且无额外条件时返回 None。
    """
    from src.config import get_settings

    conditions = [{key: value} for key, value in (extra or {}).items()]
    if get_settings().multi_tenant:
        from src.api.auth import get_current_tenant_id, get_current_user_id
        conditions.append({"tenant_id": get_current_tenant_id()})
        if owner_only:
            conditions.append({"owner_user_id": get_current_user_id()})
    if not conditions:
        return None
    if len(conditions) == 1:
        return conditions[0]
    return {"$and": conditions}


@router.get("/knowledge")
async def list_knowledge(
    category: str | None = Query(default=None),
    search: str | None = Query(default=None),
    knowledge_scope: str | None = None,
    page: int = Query(default=1, ge=1),
    page_size: int = Query(default=20, ge=10, le=100),
):
    """按当前身份列出系统、租户和个人知识条目。"""
    logger.debug(
        "知识库列表入口",
        category=category or "",
        search=search or "",
        knowledge_scope=knowledge_scope or "",
        page=page,
        page_size=page_size,
    )
    try:
        from src.memory.vector_store import get_vector_store
        from src.knowledge.governance import normalize_knowledge_scope
        from src.knowledge.retrieval import build_accessible_knowledge_filters

        store = await get_vector_store()
        filter_groups = build_accessible_knowledge_filters(category=category or "")
        if knowledge_scope:
            normalized_scope = normalize_knowledge_scope(knowledge_scope).value
            filter_groups = [
                filters for filters in filter_groups
                if filters.get("visibility") == normalized_scope
            ]
        results_by_id = {}
        for filters in filter_groups:
            scoped_results = await store.get_by_filter(filters, limit=10000)
            logger.info(
                "知识库分范围查询完成",
                visibility=filters.get("visibility", ""),
                count=len(scoped_results),
            )
            for result in scoped_results:
                results_by_id[result.id] = result
        results = list(results_by_id.values())
        logger.info(
            "知识库查询完成",
            total_ids=len(results),
            scope_count=len(filter_groups),
        )
        entries = []
        from src.api.auth import get_current_role, get_current_tenant_id, get_current_user_id
        from src.knowledge.governance import can_manage_knowledge_resource

        for result in results:
            meta = dict(result.metadata or {})
            is_user = meta.get("source") == "user_upload"
            visibility = str(meta.get("visibility", "system" if not is_user else "tenant"))
            resource_tenant_id = meta.get("tenant_id")
            if resource_tenant_id is None and visibility != "system":
                resource_tenant_id = get_current_tenant_id()
            can_delete = is_user and can_manage_knowledge_resource(
                visibility,
                role=get_current_role(),
                current_tenant_id=get_current_tenant_id(),
                resource_tenant_id=resource_tenant_id,
                current_user_id=get_current_user_id(),
                owner_user_id=meta.get("owner_user_id"),
            )
            entries.append({
                "id": result.id,
                "content": result.content,
                "category": meta.get("category", ""),
                "datasource": meta.get("datasource", ""),
                "table_name": meta.get("table_name", ""),
                "source": meta.get("source", "unknown"),
                "source_file": meta.get("source_file", ""),
                "scope": visibility,
                "tag_ids": json.loads(meta.get("tag_ids_json", "[]") or "[]"),
                "tags": [tag for tag in str(meta.get("tags", "")).split(",") if tag],
                "is_builtin": meta.get("visibility") == "system" or not is_user,
                "tenant_id": resource_tenant_id,
                "owner_user_id": meta.get("owner_user_id"),
                "can_delete": can_delete,
            })
        if search:
            q = search.lower()
            entries = [e for e in entries if q in e["content"].lower()
                       or q in e.get("table_name", "").lower()]
        total = len(entries)
        start = (page - 1) * page_size
        paged = entries[start:start + page_size]
        result = {"entries": paged, "total": total, "page": page, "page_size": page_size}
        logger.info("知识库列表完成", total=total, returned=len(paged))
        return result
    except ValueError as exc:
        logger.warning("知识库列表范围无效", error=str(exc))
        raise HTTPException(400, str(exc)) from exc
    except Exception as e:
        logger.error("知识库加载失败", error=str(e), exc_info=True)
        return {"entries": [], "total": 0, "page": page, "page_size": page_size}


# 方法作用：过滤知识文档格式并检查单文件与累计上传大小。
# Args: files - 上传文件；max_bytes - 单文件上限；max_total_bytes - 累计上限。
# Returns: 已读取文件和格式错误明细。
async def _read_knowledge_uploads(
    files: list[UploadFile],
    max_bytes: int,
    max_total_bytes: int,
) -> tuple[list[tuple[UploadFile, bytes]], list[dict]]:
    logger.debug("读取知识文档上传入口", file_count=len(files), max_bytes=max_bytes)
    supported_exts = {".md", ".txt", ".pdf", ".docx", ".doc", ".markdown", ".csv"}
    prepared: list[tuple[UploadFile, bytes]] = []
    errors: list[dict] = []
    total_bytes = 0
    for uploaded_file in files:
        if not uploaded_file.filename:
            continue
        extension = os.path.splitext(uploaded_file.filename)[1].lower()
        if extension not in supported_exts:
            errors.append({"file": uploaded_file.filename, "error": f"不支持的文件格式: {extension}"})
            continue
        content = await uploaded_file.read(max_bytes + 1)
        if len(content) > max_bytes:
            logger.warning(
                "知识文件超过大小限制",
                filename=uploaded_file.filename,
                size=len(content),
                limit=max_bytes,
            )
            raise HTTPException(413, f"文件 '{uploaded_file.filename}' 超过大小限制")
        total_bytes += len(content)
        if total_bytes > max_total_bytes:
            logger.warning("知识文件累计大小超限", total_bytes=total_bytes, limit=max_total_bytes)
            raise HTTPException(413, "上传文件累计大小超过限制")
        prepared.append((uploaded_file, content))
    logger.info("读取知识文档上传完成", accepted=len(prepared), total_bytes=total_bytes)
    return prepared, errors


# 方法作用：保存原始知识文件、创建上传任务并启动后台索引。
# Args: prepared - 已读取文件；scope - 知识范围；tag_ids - 标签 ID；tag_names - 标签名；datasource - 数据源；config - 分块配置；category - 分类。
# Returns: 已创建任务和处理错误明细。
async def _schedule_knowledge_uploads(
    prepared: list[tuple[UploadFile, bytes]],
    scope: str,
    tag_ids: list[int],
    tag_names: list[str],
    datasource: str,
    config,
    category: str,
) -> tuple[list[dict], list[dict]]:
    import asyncio
    from src.knowledge.file_store import get_file_store
    from src.knowledge.upload_manager import get_upload_manager

    logger.debug("调度知识文档上传入口", file_count=len(prepared), scope=scope)
    manager = get_upload_manager()
    tasks: list[dict] = []
    errors: list[dict] = []
    for uploaded_file, content in prepared:
        try:
            file_id = await get_file_store().save(
                uploaded_file.filename,
                content,
                knowledge_scope=scope,
                datasource=datasource,
                tag_ids=tag_ids,
            )
            task = manager.create(
                uploaded_file.filename,
                knowledge_scope=scope,
                tag_ids=tag_ids,
                tag_names=tag_names,
                datasource=datasource,
            )
            tasks.append({"task_id": task.id, "file_name": uploaded_file.filename, "file_id": file_id})
            asyncio.create_task(manager.process(task, content, config, category))
        except HTTPException:
            raise
        except Exception as exc:
            logger.error("知识文档任务调度失败", filename=uploaded_file.filename, exc_info=True)
            errors.append({"file": uploaded_file.filename, "error": str(exc)})
    logger.info("调度知识文档上传完成", tasks=len(tasks), errors=len(errors))
    return tasks, errors


# 方法作用：解析上传标签并校验当前身份可见性及公共知识标签范围。
# Args: raw_tag_ids - 逗号分隔标签；scope - 知识范围；tenant_id - 租户；user_id - 用户。
# Returns: 去重后的标签 ID 和标签名称。
async def _resolve_upload_tags(
    raw_tag_ids: str,
    scope: str,
    tenant_id: int,
    user_id: int,
) -> tuple[list[int], list[str]]:
    logger.debug("解析知识上传标签入口", scope=scope, tenant_id=tenant_id, user_id=user_id)
    try:
        requested_ids = list(dict.fromkeys(
            int(value.strip()) for value in raw_tag_ids.split(",") if value.strip()
        ))
    except ValueError as exc:
        logger.warning("知识文件上传标签格式无效", tag_ids=raw_tag_ids)
        raise HTTPException(400, "tag_ids 必须是逗号分隔的整数") from exc
    if len(requested_ids) > 20:
        raise HTTPException(400, "单个文档最多选择 20 个标签")
    if not requested_ids:
        logger.info("解析知识上传标签完成", count=0)
        return [], []
    from src.knowledge.tag_store import get_knowledge_tag_store

    tags = await get_knowledge_tag_store().get_visible_by_ids(
        requested_ids,
        tenant_id=tenant_id,
        user_id=user_id,
    )
    if {tag.id for tag in tags} != set(requested_ids):
        logger.warning("知识文件上传标签不可见", requested=requested_ids)
        raise HTTPException(400, "包含不存在、已停用或无权使用的标签")
    if scope != "private" and any(tag.scope != "global" for tag in tags):
        logger.warning("公共知识上传标签范围拒绝", knowledge_scope=scope)
        raise HTTPException(400, "系统或租户公共知识只能使用全局标签")
    names = [tag.name for tag in tags]
    logger.info("解析知识上传标签完成", count=len(requested_ids))
    return requested_ids, names


@router.post("/knowledge/docs/upload")
async def upload_knowledge_docs(
    files: list[UploadFile] = File(...),
    strategy: str = Query(default="auto"),
    chunk_size: int = Query(default=800, ge=200, le=4000),
    chunk_overlap: int = Query(default=100, ge=0, le=500),
    category: str = Query(default=""),
    knowledge_scope: str = "private",
    tag_ids: str = "",
    datasource: str = "",
):
    """按范围和标签批量上传文档并异步索引。

    system 仅超级管理员，tenant 仅租户管理员或超级管理员，private 归当前用户。
    """
    from src.config import get_settings
    from src.api.auth import (
        get_current_role, get_current_tenant_id, get_current_user_id,
    )
    from src.knowledge.governance import can_write_knowledge_scope, normalize_knowledge_scope

    settings = get_settings()
    max_upload_bytes = max(1, int(settings.max_upload_bytes))
    max_upload_files = max(1, int(getattr(settings, "max_upload_files", 20)))
    max_upload_total_bytes = max(
        max_upload_bytes,
        int(getattr(settings, "max_upload_total_bytes", 100 * 1024 * 1024)),
    )
    if len(files) > max_upload_files:
        logger.warning("知识文件上传数量超限", file_count=len(files), limit=max_upload_files)
        raise HTTPException(413, f"单次最多上传 {max_upload_files} 个文件")
    try:
        normalized_scope = normalize_knowledge_scope(knowledge_scope).value
    except ValueError as exc:
        logger.warning("知识文件上传范围无效", knowledge_scope=knowledge_scope)
        raise HTTPException(400, str(exc)) from exc
    role = get_current_role()
    user_id = get_current_user_id()
    if not can_write_knowledge_scope(
        normalized_scope,
        role=role,
        user_id=user_id,
        multi_tenant=getattr(settings, "multi_tenant", False),
    ):
        logger.warning(
            "知识文件上传权限拒绝",
            knowledge_scope=normalized_scope,
            role=role,
            user_id=user_id,
        )
        raise HTTPException(403, f"当前角色无权写入 {normalized_scope} 知识")
    logger.info(
        "知识文件上传授权完成",
        knowledge_scope=normalized_scope,
        role=role,
        tenant_id=get_current_tenant_id(),
        user_id=user_id,
    )

    requested_tag_ids, resolved_tag_names = await _resolve_upload_tags(
        tag_ids,
        normalized_scope,
        get_current_tenant_id(),
        user_id,
    )
    logger.debug(
        "知识文件上传入口",
        file_count=len(files),
        max_upload_bytes=max_upload_bytes,
        knowledge_scope=normalized_scope,
        tag_count=len(requested_tag_ids),
        datasource=datasource,
    )

    prepared_files, errors = await _read_knowledge_uploads(
        files,
        max_upload_bytes,
        max_upload_total_bytes,
    )

    from src.knowledge.doc_parser import ChunkConfig, ChunkStrategy
    try:
        strat = ChunkStrategy(strategy)
    except ValueError:
        strat = ChunkStrategy.AUTO
    config = ChunkConfig(strategy=strat, chunk_size=chunk_size, chunk_overlap=chunk_overlap)

    tasks_result, schedule_errors = await _schedule_knowledge_uploads(
        prepared_files,
        normalized_scope,
        requested_tag_ids,
        resolved_tag_names,
        datasource,
        config,
        category,
    )
    errors.extend(schedule_errors)

    result = {
        "tasks": tasks_result, "errors": errors, "total": len(tasks_result),
        "config": config.to_dict(),
        "message": "已接收文件，后台处理中。轮询 GET /api/v1/knowledge/upload/status 获取进度",
    }
    logger.info(
        "知识文件上传入口完成",
        accepted=len(tasks_result),
        errors=len(errors),
        knowledge_scope=normalized_scope,
    )
    return result


@router.get("/knowledge/upload/status")
async def upload_status(task_id: str = Query(default="")):
    """查询上传任务进度（单任务或全部最近任务）。"""
    from src.knowledge.upload_manager import get_upload_manager
    mgr = get_upload_manager()
    if task_id:
        t = mgr.get(task_id)
        if t is None:
            raise HTTPException(404, f"任务 '{task_id}' 未找到")
        return {"task": t.to_dict()}
    return {"tasks": mgr.list_recent()}


@router.get("/knowledge/tags")
async def search_knowledge_tags(
    q: str = Query(default="", max_length=128),
    include_inactive: bool = Query(default=False),
    limit: int = Query(default=30, ge=1, le=100),
):
    """搜索当前用户可见的全局标签和个人标签。"""
    from src.api.auth import (
        get_current_role, get_current_tenant_id, get_current_user_id,
    )
    from src.knowledge.governance import is_super_admin
    from src.knowledge.tag_store import get_knowledge_tag_store

    role = get_current_role()
    effective_include_inactive = include_inactive and is_super_admin(role)
    logger.debug(
        "知识标签搜索 API 入口",
        query=q,
        include_inactive=effective_include_inactive,
        limit=limit,
        role=role,
    )
    try:
        tags = await get_knowledge_tag_store().search(
            q,
            tenant_id=get_current_tenant_id(),
            user_id=get_current_user_id(),
            limit=limit,
            include_inactive=effective_include_inactive,
            include_all_private=is_super_admin(role),
        )
        result = {"tags": [tag.model_dump() for tag in tags], "total": len(tags)}
        logger.info("知识标签搜索 API 完成", count=len(tags), role=role)
        return result
    except Exception as exc:
        logger.error("知识标签搜索 API 失败", error=str(exc), exc_info=True)
        raise HTTPException(500, "知识标签搜索失败") from exc


@router.post("/knowledge/tags", status_code=201)
async def create_knowledge_tag(request: KnowledgeTagCreateRequest):
    """为当前用户创建默认仅本人可见的自定义标签。"""
    from src.api.auth import get_current_tenant_id, get_current_user_id
    from src.knowledge.tag_store import get_knowledge_tag_store

    tenant_id = get_current_tenant_id()
    user_id = get_current_user_id()
    logger.debug("创建个人知识标签 API 入口", name=request.name, tenant_id=tenant_id, user_id=user_id)
    try:
        tag = await get_knowledge_tag_store().create_personal(
            request.name,
            tenant_id=tenant_id,
            user_id=user_id,
            description=request.description,
            aliases=request.aliases,
        )
        result = tag.model_dump()
        logger.info("创建个人知识标签 API 完成", tag_id=result["id"], user_id=user_id)
        return result
    except ValueError as exc:
        logger.warning("创建个人知识标签 API 输入拒绝", error=str(exc), user_id=user_id)
        raise HTTPException(400, str(exc)) from exc
    except Exception as exc:
        logger.error("创建个人知识标签 API 失败", error=str(exc), exc_info=True)
        raise HTTPException(500, "创建个人标签失败") from exc


@router.post("/knowledge/tags/global", status_code=201)
async def create_global_knowledge_tag(request: KnowledgeTagCreateRequest):
    """由超级管理员创建或重新启用平台全局标签。"""
    from src.api.auth import get_current_role, require_super_admin
    from src.knowledge.tag_store import get_knowledge_tag_store

    logger.debug("创建全局知识标签 API 入口", name=request.name, role=get_current_role())
    require_super_admin()
    try:
        tag = await get_knowledge_tag_store().create_global(
            request.name,
            actor_role=get_current_role(),
            tag_group=request.tag_group,
            description=request.description,
            aliases=request.aliases,
        )
        result = tag.model_dump()
        logger.info("创建全局知识标签 API 完成", tag_id=result["id"])
        return result
    except ValueError as exc:
        logger.warning("创建全局知识标签 API 输入拒绝", error=str(exc))
        raise HTTPException(400, str(exc)) from exc
    except Exception as exc:
        logger.error("创建全局知识标签 API 失败", error=str(exc), exc_info=True)
        raise HTTPException(500, "创建全局标签失败") from exc


@router.post("/knowledge/tags/{tag_id}/promote")
async def promote_knowledge_tag(tag_id: int):
    """由超级管理员把个人标签提升为平台全局标签。"""
    from src.api.auth import get_current_role, require_super_admin
    from src.knowledge.tag_store import get_knowledge_tag_store

    logger.debug("提升知识标签 API 入口", tag_id=tag_id, role=get_current_role())
    require_super_admin()
    try:
        tag = await get_knowledge_tag_store().promote_to_global(
            tag_id,
            actor_role=get_current_role(),
        )
        result = tag.model_dump()
        logger.info("提升知识标签 API 完成", source_tag_id=tag_id, global_tag_id=result["id"])
        return result
    except ValueError as exc:
        logger.warning("提升知识标签 API 未命中", tag_id=tag_id, error=str(exc))
        raise HTTPException(404, str(exc)) from exc
    except Exception as exc:
        logger.error("提升知识标签 API 失败", tag_id=tag_id, error=str(exc), exc_info=True)
        raise HTTPException(500, "提升标签失败") from exc


@router.patch("/knowledge/tags/{tag_id}")
async def update_knowledge_tag_status(tag_id: int, request: KnowledgeTagStatusRequest):
    """启用或停用全局标签或当前用户自己的个人标签。"""
    from src.api.auth import (
        get_current_role, get_current_tenant_id, get_current_user_id,
    )
    from src.knowledge.tag_store import get_knowledge_tag_store

    logger.debug(
        "更新知识标签状态 API 入口",
        tag_id=tag_id,
        is_active=request.is_active,
        role=get_current_role(),
    )
    try:
        updated = await get_knowledge_tag_store().set_active(
            tag_id,
            request.is_active,
            actor_role=get_current_role(),
            tenant_id=get_current_tenant_id(),
            user_id=get_current_user_id(),
        )
        if not updated:
            raise HTTPException(404, f"标签 '{tag_id}' 不存在或无权管理")
        logger.info("更新知识标签状态 API 完成", tag_id=tag_id, is_active=request.is_active)
        return {"status": "ok", "id": tag_id, "is_active": request.is_active}
    except HTTPException:
        raise
    except Exception as exc:
        logger.error("更新知识标签状态 API 失败", tag_id=tag_id, error=str(exc), exc_info=True)
        raise HTTPException(500, "更新标签状态失败") from exc


@router.get("/knowledge/test-search")
async def test_knowledge_search(q: str = Query(default=""), datasource: str = Query(default="")):
    """测试知识库检索效果——验证上传的文档是否能被检索到。

    有 q 时做语义搜索，返回匹配条目和相关性分数。
    无 q 时返回全部条目列表。
    """
    try:
        from src.memory.vector_store import get_vector_store
        from src.knowledge.retrieval import build_knowledge_filters, search_knowledge
        store = await get_vector_store()
        from src.api.auth import get_current_tenant_id
        if q:
            evidence = await search_knowledge(store, q, datasource=datasource, top_k=10)
            items = [{"rank": i + 1, "id": e.source_id, "content": e.content[:200],
                      "relevance": e.scores.get("relevance", 0.0), "citation": e.locator}
                     for i, e in enumerate(evidence)]
            return {"query": q, "results": items, "total": len(items)}
        else:
            filters = build_knowledge_filters(datasource=datasource)
            results = await store.get_by_filter(filters, limit=1000)
            return {"total": len(results), "ids": [r.id for r in results]}
    except Exception as e:
        raise HTTPException(500, f"知识库检索测试失败: {e}")


@router.get("/knowledge/docs")
async def list_knowledge_docs():
    """列出当前身份可见文档并附带服务端删除权限。"""
    from src.api.auth import get_current_role, get_current_tenant_id, get_current_user_id
    from src.knowledge.governance import can_manage_knowledge_resource
    from src.knowledge.file_store import get_file_store

    logger.debug("知识文档列表 API 入口", role=get_current_role())
    docs = await get_file_store().list_files()
    for doc in docs:
        scope = str(doc.get("scope", "system" if doc.get("is_builtin") else "private"))
        doc["can_delete"] = not doc.get("is_builtin", False) and can_manage_knowledge_resource(
            scope,
            role=get_current_role(),
            current_tenant_id=get_current_tenant_id(),
            resource_tenant_id=doc.get("tenant_id", get_current_tenant_id()),
            current_user_id=get_current_user_id(),
            owner_user_id=doc.get("owner_user_id"),
        )
    result = {"docs": docs, "total": len(docs)}
    logger.info("知识文档列表 API 完成", total=len(docs), role=get_current_role())
    return result


@router.get("/knowledge/docs/{doc_name}/content")
async def get_doc_content(doc_name: str, knowledge_scope: str = ""):
    """获取已索引文档的内容（从 PG 读取，回退磁盘）。"""
    from src.knowledge.file_store import get_file_store
    store = get_file_store()
    doc = await store.get_by_name(doc_name, knowledge_scope=knowledge_scope)
    if doc:
        raw = doc["file_data"]
        size = doc["size"]
    else:
        if knowledge_scope and knowledge_scope != "system":
            raise HTTPException(404, f"文档 '{doc_name}' 未找到")
        # 磁盘回退仅用于系统内置目录。
        base_dir = os.path.realpath(os.path.join(os.path.dirname(__file__), "..", "..", "docs", "metrics"))
        doc_path = os.path.realpath(os.path.join(base_dir, os.path.basename(doc_name)))
        if not doc_path.startswith(base_dir):
            raise HTTPException(403, "禁止访问")
        if not os.path.isfile(doc_path):
            raise HTTPException(404, f"文档 '{doc_name}' 未找到")
        with open(doc_path, "rb") as fp:
            raw = fp.read()
        size = os.path.getsize(doc_path)

    ext = os.path.splitext(doc_name)[1].lower()
    result = {"name": doc_name, "size": size, "ext": ext, "type": "text", "content": "", "raw_url": ""}
    if ext == ".pdf":
        result["type"] = "pdf"
        result["raw_url"] = (
            f"/api/v1/knowledge/docs/{doc_name}/raw"
            f"?knowledge_scope={knowledge_scope or 'system'}"
        )
        from src.knowledge.doc_parser import extract_text
        result["content"] = extract_text(doc_name, raw)
    elif ext in (".docx", ".doc"):
        result["type"] = "word"
        result["content"] = _docx_to_html(raw)
    else:
        result["content"] = raw.decode("utf-8", errors="replace")
    return result


@router.get("/knowledge/docs/{doc_name}/raw")
async def get_doc_raw(doc_name: str, knowledge_scope: str = ""):
    """返回原始文件（从 PG 读取，回退磁盘，用于 PDF iframe 渲染）。"""
    from fastapi.responses import Response
    from src.knowledge.file_store import get_file_store
    doc = await get_file_store().get_by_name(doc_name, knowledge_scope=knowledge_scope)
    if doc:
        return Response(content=bytes(doc["file_data"]), media_type=doc["content_type"])
    if knowledge_scope and knowledge_scope != "system":
        raise HTTPException(404, f"文档 '{doc_name}' 未找到")
    # 磁盘回退仅用于系统内置目录。
    from fastapi.responses import FileResponse
    base_dir = os.path.realpath(os.path.join(os.path.dirname(__file__), "..", "..", "docs", "metrics"))
    doc_path = os.path.realpath(os.path.join(base_dir, os.path.basename(doc_name)))
    if not doc_path.startswith(base_dir):
        raise HTTPException(403, "禁止访问")
    if not os.path.isfile(doc_path):
        raise HTTPException(404, f"文档 '{doc_name}' 未找到")
    ext = os.path.splitext(doc_name)[1].lower()
    media_map = {".pdf": "application/pdf", ".txt": "text/plain", ".md": "text/markdown"}
    return FileResponse(doc_path, media_type=media_map.get(ext, "application/octet-stream"))


def _docx_to_html(raw: bytes) -> str:
    """将 Word 文档转换为转义后的简单 HTML。

    Args:
        raw: Word 文档二进制内容。

    Returns:
        可安全嵌入前端的 HTML 文本。
    """
    logger.debug("Word HTML 转换入口", size=len(raw))
    try:
        from docx import Document
        doc = Document(io.BytesIO(raw))
        parts = ['<div style="font-family: sans-serif; line-height: 1.8;">']
        for para in doc.paragraphs:
            text = html.escape(para.text.strip(), quote=True)
            if not text:
                parts.append('<br/>')
                continue
            if para.style and para.style.name and para.style.name.startswith("Heading"):
                level = para.style.name.replace("Heading", "").strip()
                try:
                    lv = max(1, min(6, int(level)))
                except ValueError:
                    lv = 2
                parts.append(f'<h{lv} style="margin:12px 0 6px;">{text}</h{lv}>')
            else:
                parts.append(f'<p style="margin:4px 0;">{text}</p>')
        # 处理表格
        for table in doc.tables:
            parts.append('<table border="1" cellpadding="6" cellspacing="0" style="border-collapse:collapse; width:100%; margin:8px 0;">')
            for row in table.rows:
                parts.append('<tr>')
                for cell in row.cells:
                    cell_text = html.escape(cell.text, quote=True)
                    parts.append(f'<td style="padding:4px 8px;">{cell_text}</td>')
                parts.append('</tr>')
            parts.append('</table>')
        parts.append('</div>')
        result = "\n".join(parts)
        logger.info("Word HTML 转换完成", size=len(raw), output_size=len(result))
        return result
    except Exception as exc:
        logger.error("Word HTML 转换失败", error=str(exc), exc_info=True)
        return "<p>无法渲染 Word 文档内容</p>"


@router.delete("/knowledge/{entry_id}")
async def delete_knowledge_entry(entry_id: str):
    """按知识范围和管理员角色删除指定用户上传条目。"""
    logger.debug("删除知识条目 API 入口", entry_id=entry_id)
    try:
        from src.memory.vector_store import get_vector_store
        store = await get_vector_store()
        entry = await store.get_by_id(entry_id)
        if entry is None:
            logger.warning("删除知识条目未命中", entry_id=entry_id)
            raise HTTPException(404, f"知识条目 '{entry_id}' 未找到")
        meta = dict(entry.metadata or {})
        if meta.get("source") != "user_upload":
            logger.warning("删除知识条目拒绝", entry_id=entry_id, reason="非用户上传")
            raise HTTPException(403, "目录扫描或系统内置知识条目不可通过 API 删除")
        from src.api.auth import (
            get_current_role, get_current_tenant_id, get_current_user_id,
        )
        from src.knowledge.governance import can_manage_knowledge_resource

        visibility = str(meta.get("visibility") or "private")
        current_tenant_id = get_current_tenant_id()
        resource_tenant_id = meta.get("tenant_id")
        if resource_tenant_id is None and visibility != "system":
            resource_tenant_id = current_tenant_id
        owner_user_id = meta.get("owner_user_id")
        allowed = can_manage_knowledge_resource(
            visibility,
            role=get_current_role(),
            current_tenant_id=current_tenant_id,
            resource_tenant_id=resource_tenant_id,
            current_user_id=get_current_user_id(),
            owner_user_id=owner_user_id,
        )
        if not allowed:
            logger.warning(
                "删除知识条目权限拒绝",
                entry_id=entry_id,
                visibility=visibility,
                role=get_current_role(),
            )
            raise HTTPException(403, "无权删除该知识条目")
        await store.delete_by_ids([entry_id])
        result = {"status": "ok", "id": entry_id}
        logger.info("删除知识条目 API 完成", entry_id=entry_id, visibility=visibility)
        return result
    except HTTPException:
        raise
    except Exception as e:
        logger.error("删除知识条目 API 失败", entry_id=entry_id, error=str(e), exc_info=True)
        raise HTTPException(500, str(e))


@router.delete("/knowledge/docs/{doc_name}")
async def delete_knowledge_doc(doc_name: str, knowledge_scope: str = ""):
    """按范围删除当前身份可管理的原文件及全部关联向量。"""
    logger.debug(
        "删除知识文档 API 入口",
        doc_name=doc_name,
        knowledge_scope=knowledge_scope,
    )
    from src.memory.vector_store import get_vector_store
    from src.api.auth import get_current_role, get_current_tenant_id, get_current_user_id
    from src.knowledge.governance import can_manage_knowledge_resource, normalize_knowledge_scope
    from src.knowledge.retrieval import build_accessible_knowledge_filters

    if knowledge_scope:
        try:
            normalized_scope = normalize_knowledge_scope(knowledge_scope).value
        except ValueError as exc:
            raise HTTPException(400, str(exc)) from exc
    else:
        normalized_scope = ""
    store = await get_vector_store()
    results = []
    for filters in build_accessible_knowledge_filters():
        scoped_filters = {**filters, "source_file": doc_name}
        results.extend(await store.get_by_filter(scoped_filters, limit=10000))
    manageable = []
    for entry in results:
        metadata = dict(entry.metadata or {})
        visibility = str(metadata.get("visibility") or "private")
        if normalized_scope and visibility != normalized_scope:
            continue
        if metadata.get("source") != "user_upload":
            logger.warning("删除知识文档拒绝", doc_name=doc_name, reason="系统目录或内置文档")
            raise HTTPException(403, f"内置文档 '{doc_name}' 不可删除")
        resource_tenant_id = metadata.get("tenant_id")
        if resource_tenant_id is None and visibility != "system":
            resource_tenant_id = get_current_tenant_id()
        if can_manage_knowledge_resource(
            visibility,
            role=get_current_role(),
            current_tenant_id=get_current_tenant_id(),
            resource_tenant_id=resource_tenant_id,
            current_user_id=get_current_user_id(),
            owner_user_id=metadata.get("owner_user_id"),
        ):
            manageable.append(entry)
    if results and not manageable:
        logger.warning("删除知识文档权限拒绝", doc_name=doc_name, role=get_current_role())
        raise HTTPException(403, "无权删除该知识文档")
    # docs/metrics 仅包含内置文档，禁止通过管理 API 删除。
    base_dir = os.path.realpath(os.path.join(os.path.dirname(__file__), "..", "..", "docs", "metrics"))
    doc_path = os.path.realpath(os.path.join(base_dir, os.path.basename(doc_name)))
    if not doc_path.startswith(base_dir):
        raise HTTPException(403, "禁止访问")
    if os.path.isfile(doc_path):
        raise HTTPException(403, f"内置文档 '{doc_name}' 不可删除")
    # 从 PG 删除原始文件和 VectorStore 关联条目。
    from src.knowledge.file_store import get_file_store
    deleted_file = await get_file_store().delete(doc_name, knowledge_scope=normalized_scope)
    if manageable:
        await store.delete_by_ids([entry.id for entry in manageable])
    if not deleted_file and not manageable:
        raise HTTPException(404, f"文档 '{doc_name}' 未找到")
    result = {"status": "ok", "doc": doc_name}
    logger.info(
        "删除知识文档 API 完成",
        doc_name=doc_name,
        knowledge_scope=normalized_scope,
        vector_count=len(manageable),
        file_deleted=deleted_file,
    )
    return result
