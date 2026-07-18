"""11.1 + 2.3.7-9 API 路由 — chat / schema / datasources / health (13 端点)。"""

from __future__ import annotations

import io
import html
import json
import os
import time
import uuid

from fastapi import APIRouter, Body, File, HTTPException, Query, UploadFile

from src.api.schemas import (
    ChatRequest, ChatResponse, ColumnCommentRequest,
    DataSourceCreateRequest, DataSourceInfo, HealthResponse, KnowledgeTagCreateRequest,
    KnowledgeTagStatusRequest, MCPServerCreate, TableInfo,
)
from src.exceptions import DataSourceNotFoundError
from src.llm.client import is_llm_available
from src.logging_config import get_logger

logger = get_logger(__name__)
router = APIRouter()
_started_at = time.time()


def _app():
    from src.graph.workflow import app
    return app


def _registry():
    from src.datasource.registry import get_registry
    return get_registry()


def _schema_manager():
    """获取全局 SchemaManager 实例。

    Returns:
        SchemaManager 单例。
    """
    from src.knowledge.schema_manager import get_schema_manager
    return get_schema_manager()


def _knowledge_where(extra: dict | None = None, owner_only: bool = False) -> dict | None:
    """构建兼容 ChromaDB 的知识库租户过滤条件。

    Args:
        extra: 额外 metadata 精确过滤条件。
        owner_only: 是否同时限制为当前用户创建。

    Returns:
        ChromaDB where 条件；单租户且无额外条件时返回 None。
    """
    from src.config import get_settings

    conditions = [{key: value} for key, value in (extra or {}).items()]
    if get_settings().multi_tenant:
        from src.api.auth import get_current_tenant_id, get_current_user_id
        conditions.append({"tenant_id": get_current_tenant_id()})
        if owner_only:
            conditions.append({"owner_user_id": get_current_user_id()})
    if not conditions:
        return None
    if len(conditions) == 1:
        return conditions[0]
    return {"$and": conditions}


# ---- Chat (11.1.1-2) ----

@router.post("/chat")
async def chat(req: ChatRequest):
    """统一 chat 端点：stream=False 返回 JSON，stream=True 返回 SSE 流式。"""
    selected_datasources = (
        req.datasources if req.datasources and len(req.datasources) > 1
        else [req.datasource]
    )
    logger.debug(
        "Chat 请求入口",
        datasource=req.datasource,
        selected_count=len(selected_datasources),
        stream=req.stream,
    )
    if req.stream:
        from fastapi.responses import StreamingResponse
        from src.api.streaming import stream_analysis
        from src.api.auth import get_current_role, get_current_tenant_id, get_current_user_id
        logger.info(
            "Chat 流式响应已创建",
            datasource=req.datasource,
            selected_count=len(selected_datasources),
        )
        return StreamingResponse(
            stream_analysis(
                req.query, req.datasource, req.session_id or "", req.datasources,
                tenant_id=get_current_tenant_id(),
                user_id=get_current_user_id(),
                user_role=get_current_role(),
            ),
            media_type="text/event-stream",
            headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
        )
    import uuid as _uuid
    sid = req.session_id or str(_uuid.uuid4())
    # 非流式也保存会话元数据
    import asyncio as _asyncio
    try:
        from src.memory.session_store import get_session_store
        if req.session_id:
            _asyncio.create_task(get_session_store().touch(sid, req.datasource, req.query))
        else:
            _asyncio.create_task(get_session_store().create(sid, req.datasource, req.query))
    except Exception:
        pass
    from src.api.auth import scope_thread_id
    from src.api.auth import get_current_role, get_current_tenant_id, get_current_user_id
    cfg = {"configurable": {"thread_id": scope_thread_id(sid)}}
    result = await _app().ainvoke({
        "user_query": req.query,
        "datasource": req.datasource,
        "session_id": sid,
        "selected_datasources": selected_datasources,
        "allowed_columns": [],
        "row_filter_sql": "",
        "tenant_id": get_current_tenant_id(),
        "user_id": get_current_user_id(),
        "user_role": get_current_role(),
    }, cfg)
    f = result.get("final_response", {})
    return ChatResponse(
        success=f.get("success", True), session_id=sid[:8],
        user_query=req.query,
        sql=f.get("sql", result.get("generated_sql", "")),
        sql_statements=f.get("sql_statements", []),
        data=f.get("data", result.get("query_result_sample", [])),
        row_count=f.get("row_count", result.get("query_result_full_count", 0)),
        truncated=bool(f.get("truncated", result.get("query_result_truncated", False))),
        analysis=f.get("analysis", result.get("analysis_result", {})),
        chart=f.get("chart", result.get("chart_config", {})),
    )


@router.post("/chat/stream")
async def chat_stream(req: ChatRequest):
    """（保留向后兼容）独立流式端点，等价于 /chat + stream=True。"""
    logger.debug(
        "兼容流式 Chat 入口",
        datasource=req.datasource,
        selected_count=len(req.datasources or [req.datasource]),
    )
    from fastapi.responses import StreamingResponse
    from src.api.streaming import stream_analysis
    from src.api.auth import get_current_role, get_current_tenant_id, get_current_user_id
    logger.info("兼容流式 Chat 响应已创建", datasource=req.datasource)
    return StreamingResponse(
        stream_analysis(
            req.query,
            req.datasource,
            req.session_id or "",
            req.datasources,
            tenant_id=get_current_tenant_id(),
            user_id=get_current_user_id(),
            user_role=get_current_role(),
        ),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


# ---- Schema (11.1.3-6) ----

@router.get("/schema/tables")
async def list_tables(
    datasource: str = Query(default="demo"),
    page: int = Query(default=1, ge=1), page_size: int = Query(default=20, ge=1, le=100),
    search: str = Query(default=""),
) -> dict:
    try:
        ds = await _registry().resolve(datasource)
    except DataSourceNotFoundError:
        raise HTTPException(404, f"数据源 '{datasource}' 未找到")
    if ds.schema is None:
        logger.info("Schema 路由触发延迟内省", datasource=datasource)
        ds.schema = await _schema_manager().get_or_fetch_schema(datasource)
    tables = []
    for t in (ds.schema.tables if ds.schema else []):
        if search and search.lower() not in t.name.lower():
            continue
        tables.append(TableInfo(name=t.name, description=t.description,
            columns=[{"name": c.name, "type": c.type, "comment": c.comment,
                      "is_indexed": c.is_indexed, "is_primary_key": c.is_primary_key}
                     for c in t.columns],
            row_count_estimate=t.row_count_estimate))
    total = len(tables)
    start = (page - 1) * page_size
    return {"tables": tables[start:start+page_size], "datasource": datasource,
            "total": total, "page": page, "page_size": page_size}


@router.get("/schema/tables/{table_name}")
async def get_table(table_name: str, datasource: str = Query(default="demo")):
    try:
        ds = await _registry().resolve(datasource)
    except DataSourceNotFoundError:
        raise HTTPException(404, f"数据源 '{datasource}' 未找到")
    if ds.schema is None:
        logger.info("表详情路由触发延迟内省", datasource=datasource)
        ds.schema = await _schema_manager().get_or_fetch_schema(datasource)
    for t in (ds.schema.tables if ds.schema else []):
        if t.name == table_name:
            return TableInfo(name=t.name, description=t.description,
                columns=[{"name": c.name, "type": c.type, "comment": c.comment,
                          "is_nullable": c.is_nullable, "is_primary_key": c.is_primary_key}
                         for c in t.columns],
                row_count_estimate=t.row_count_estimate)
    raise HTTPException(404, f"表 '{table_name}' 未找到")


@router.post("/schema/refresh")
async def refresh_schema(datasource: str = Query(default="demo")):
    """验证数据源后真实刷新 Schema 缓存。

    Args:
        datasource: 数据源名称。

    Returns:
        刷新状态和表数量。
    """
    logger.debug("Schema 刷新路由入口", datasource=datasource)
    try:
        await _registry().resolve(datasource)
    except DataSourceNotFoundError:
        logger.warning("Schema 刷新数据源不存在", datasource=datasource)
        raise HTTPException(404, f"数据源 '{datasource}' 未找到")
    snapshot = await _schema_manager().refresh(datasource)
    result = {
        "status": "ok",
        "message": "刷新完成",
        "datasource": datasource,
        "table_count": len(snapshot.tables),
    }
    logger.info("Schema 刷新路由完成", datasource=datasource, table_count=result["table_count"])
    return result


@router.put("/schema/tables/{table_name}/columns/{column_name}/comment")
async def update_column_comment(
    table_name: str, column_name: str, req: ColumnCommentRequest,
    datasource: str = Query(default="demo"),
):
    """更新指定数据源中的字段备注。

    Args:
        table_name: 表名。
        column_name: 字段名。
        req: 备注请求体。
        datasource: 数据源名称。

    Returns:
        更新后的字段备注摘要。
    """
    logger.debug("字段备注路由入口", datasource=datasource, table=table_name, column=column_name)
    try:
        await _registry().resolve(datasource)
    except DataSourceNotFoundError:
        raise HTTPException(404, f"数据源 '{datasource}' 未找到")
    updated = await _schema_manager().update_column_comment(
        datasource, table_name, column_name, req.comment,
    )
    if not updated:
        raise HTTPException(404, f"字段 '{table_name}.{column_name}' 未找到")
    logger.info("字段备注路由完成", datasource=datasource, table=table_name, column=column_name)
    return {"status": "ok", "table": table_name, "column": column_name, "comment": req.comment}


# ---- 数据源管理 (2.3.7-9) ----

@router.post("/datasources", status_code=201)
async def register_datasource(req: DataSourceCreateRequest):
    """把外部数据源注册到全局 Provider/Registry。

    Args:
        req: 数据源注册请求体。

    Returns:
        已注册数据源摘要。
    """
    from src.datasource.providers.external import ExternalDataSourceProvider
    registry = _registry()
    provider = registry.get_provider("external")
    if provider is None:
        provider = ExternalDataSourceProvider()
        registry.register_provider("external", provider)
    ds = await provider.register(req)
    registry.invalidate(ds.name)
    logger.info("数据源注册路由完成", datasource=ds.name)
    return DataSourceInfo(name=ds.name, dialect=ds.dialect, version=ds.version,
                          mode=ds.mode, host=ds.host, database=ds.database,
                          description=ds.description)


@router.delete("/datasources/{name}")
async def delete_datasource(name: str):
    """从全局 Registry 删除数据源并释放连接。

    Args:
        name: 数据源名称。

    Returns:
        删除状态。
    """
    logger.debug("数据源删除路由入口", datasource=name)
    if not await _registry().unregister(name):
        logger.warning("数据源删除目标不存在", datasource=name)
        raise HTTPException(404, f"数据源 '{name}' 未找到")
    logger.info("数据源删除路由完成", datasource=name)
    return {"status": "ok", "message": f"数据源 '{name}' 已删除"}


@router.get("/datasources")
async def list_datasources(page: int = Query(default=1, ge=1), page_size: int = Query(default=20, ge=1, le=100)):
    items = await _registry().list_all()
    total = len(items)
    start = (page - 1) * page_size
    return {"datasources": items[start:start+page_size], "total": total, "page": page, "page_size": page_size}


# ---- MCP Server 管理 ----


# 方法作用：规范化扩展资源作用域并校验当前身份写权限。
# Args: scope - system/tenant/private 作用域。
# Returns: 规范化作用域、当前租户、当前用户和当前角色。
def _authorize_extension_scope(scope: str) -> tuple[str, int, int, str]:
    """Skill 和 MCP 共用与知识库一致的三级写权限。"""
    logger.debug("扩展资源作用域授权入口", scope=scope)
    from src.api.auth import (
        get_current_role, get_current_tenant_id, get_current_user_id,
    )
    from src.config import get_settings
    from src.knowledge.governance import can_write_knowledge_scope, normalize_knowledge_scope

    try:
        normalized = normalize_knowledge_scope(scope).value
    except ValueError as exc:
        logger.warning("扩展资源作用域无效", scope=scope)
        raise HTTPException(400, str(exc)) from exc
    tenant_id = get_current_tenant_id()
    user_id = get_current_user_id()
    role = get_current_role()
    if not can_write_knowledge_scope(
        normalized,
        role=role,
        user_id=user_id,
        multi_tenant=get_settings().multi_tenant,
    ):
        logger.warning(
            "扩展资源作用域授权拒绝",
            scope=normalized,
            role=role,
            tenant_id=tenant_id,
            user_id=user_id,
        )
        raise HTTPException(403, f"当前角色无权写入 {normalized} 扩展资源")
    logger.info(
        "扩展资源作用域授权完成",
        scope=normalized,
        role=role,
        tenant_id=tenant_id,
        user_id=user_id,
    )
    return normalized, tenant_id, user_id, role


# 方法作用：把作用域和当前身份转换为 MCP 数据库存储所有者字段。
# Args: scope - 已规范化作用域；tenant_id - 当前租户；user_id - 当前用户。
# Returns: 数据库 tenant_id 与 owner_user_id。
def _mcp_owner_fields(scope: str, tenant_id: int, user_id: int) -> tuple[int | None, int]:
    """system 不绑定租户，tenant 不绑定个人，private 同时绑定租户和用户。"""
    logger.debug(
        "计算 MCP 所有者字段入口", scope=scope, tenant_id=tenant_id, user_id=user_id,
    )
    if scope == "system":
        result = (None, 0)
    elif scope == "tenant":
        result = (tenant_id, 0)
    else:
        result = (tenant_id, user_id)
    logger.info(
        "计算 MCP 所有者字段完成",
        scope=scope,
        resource_tenant_id=result[0],
        owner_user_id=result[1],
    )
    return result


# 方法作用：创建已注入当前认证身份的 asyncpg 连接供 MCP RLS 查询使用。
# Args: 无。
# Returns: 设置 tenant_id/user_id/role 后的 asyncpg 连接。
async def _connect_scoped_mcp_db():
    """所有 MCP 管理 SQL 必须经过连接级 RLS 身份注入。"""
    logger.debug("连接 MCP 作用域数据库入口")
    import asyncpg
    from src.api.auth import get_current_role, get_current_tenant_id, get_current_user_id
    from src.config import get_settings

    url = get_settings().database_url.replace("postgresql+asyncpg://", "postgresql://")
    connection = await asyncpg.connect(url)
    try:
        await connection.execute(
            "SELECT set_config('app.current_tenant_id', $1, false), "
            "set_config('app.current_user_id', $2, false), "
            "set_config('app.current_role', $3, false)",
            str(get_current_tenant_id()), str(get_current_user_id()), get_current_role(),
        )
    except Exception:
        await connection.close()
        logger.error("连接 MCP 作用域数据库失败", exc_info=True)
        raise
    logger.info("连接 MCP 作用域数据库完成")
    return connection


@router.get("/mcp/servers")
# 方法作用：列出当前身份可见的 MCP Server。
# Args: scope - 可选 system/tenant/private 过滤范围。
# Returns: MCP Server 列表和总数。
async def list_mcp_servers(scope: str | None = None):
    """列出当前身份可见的 system/tenant/private MCP Server。"""
    logger.debug("MCP Server 列表入口", scope=scope or "")
    from src.api.auth import get_current_role, get_current_tenant_id, get_current_user_id
    from src.knowledge.governance import is_super_admin, normalize_knowledge_scope

    tenant_id = get_current_tenant_id()
    user_id = get_current_user_id()
    role = get_current_role()
    normalized_scope = None
    if scope:
        try:
            normalized_scope = normalize_knowledge_scope(scope).value
        except ValueError as exc:
            raise HTTPException(400, str(exc)) from exc
    from src.mcp_client.client_manager import get_mcp_client_manager
    runtime_system = (
        get_mcp_client_manager().list_system_servers()
        if normalized_scope in (None, "system") else []
    )
    try:
        conn = await _connect_scoped_mcp_db()
        try:
            if is_super_admin(role):
                rows = await conn.fetch(
                    "SELECT name, scope, tenant_id, owner_user_id, transport, command, args, "
                    "url, description, is_builtin, enabled FROM mcp_servers "
                    "ORDER BY scope, name",
                )
            else:
                rows = await conn.fetch(
                    "SELECT name, scope, tenant_id, owner_user_id, transport, command, args, "
                    "url, description, is_builtin, enabled FROM mcp_servers WHERE "
                    "scope='system' OR (scope='tenant' AND tenant_id=$1) OR "
                    "(scope='private' AND tenant_id=$1 AND owner_user_id=$2) "
                    "ORDER BY scope, name",
                    tenant_id, user_id,
                )
        finally:
            await conn.close()
        servers = list(runtime_system)
        for row in rows:
            row_scope = str(row["scope"] or "tenant")
            if normalized_scope and row_scope != normalized_scope:
                continue
            servers.append({
                "name": row["name"], "scope": row_scope,
                "tenant_id": int(row["tenant_id"] or 0),
                "owner_user_id": int(row["owner_user_id"] or 0),
                "transport": row["transport"], "command": row["command"],
                "args": row["args"], "url": row["url"],
                "description": row["description"], "is_builtin": row["is_builtin"],
                "enabled": row["enabled"],
            })
        result = {"servers": servers, "total": len(servers)}
        logger.info(
            "MCP Server 列表完成", tenant_id=tenant_id, user_id=user_id, total=len(servers),
        )
        return result
    except Exception as exc:
        logger.error("MCP Server 数据库列表失败", error=str(exc), exc_info=True)
        return {"servers": runtime_system, "total": len(runtime_system)}


@router.post("/mcp/servers", status_code=201)
# 方法作用：按认证身份创建或更新指定作用域 MCP Server。
# Args: req - MCP Server 连接配置和目标作用域。
# Returns: 创建状态、作用域和运行时加载数量。
async def create_mcp_server(req: MCPServerCreate):
    """按当前身份创建或更新 system/tenant/private MCP Server。"""
    logger.debug("创建 MCP Server 入口", name=req.name, scope=req.scope)
    normalized_scope, tenant_id, user_id, _ = _authorize_extension_scope(req.scope)
    resource_tenant_id, owner_user_id = _mcp_owner_fields(
        normalized_scope, tenant_id, user_id,
    )
    try:
        import json
        conn = await _connect_scoped_mcp_db()
        try:
            existing = await conn.fetchval(
                "SELECT id FROM mcp_servers WHERE name=$1 AND scope=$2 "
                "AND COALESCE(tenant_id, 0)=COALESCE($3::int, 0) AND owner_user_id=$4",
                req.name, normalized_scope, resource_tenant_id, owner_user_id,
            )
            if existing:
                await conn.execute(
                    "UPDATE mcp_servers SET transport=$1, command=$2, args=$3, url=$4, "
                    "env_vars=$5, description=$6, enabled=$7 WHERE id=$8",
                    req.transport, req.command, req.args, req.url,
                    json.dumps(req.env_vars), req.description, req.enabled, existing,
                )
            else:
                await conn.execute(
                    "INSERT INTO mcp_servers (name, scope, tenant_id, owner_user_id, transport, "
                    "command, args, url, env_vars, description, enabled) "
                    "VALUES ($1,$2,$3,$4,$5,$6,$7,$8,$9,$10,$11)",
                    req.name, normalized_scope, resource_tenant_id, owner_user_id,
                    req.transport, req.command, req.args, req.url,
                    json.dumps(req.env_vars), req.description, req.enabled,
                )
        finally:
            await conn.close()
        from src.mcp_client.client_manager import get_mcp_client_manager
        loaded = await get_mcp_client_manager().ensure_scoped_servers(
            tenant_id, user_id, force=True,
        )
        result = {
            "status": "ok", "name": req.name, "scope": normalized_scope,
            "runtime_loaded": loaded,
        }
        logger.info(
            "创建 MCP Server 完成", name=req.name, scope=normalized_scope,
            tenant_id=tenant_id, user_id=user_id,
        )
        return result
    except HTTPException:
        raise
    except Exception as exc:
        logger.error("创建 MCP Server 失败", name=req.name, error=str(exc), exc_info=True)
        raise HTTPException(500, f"创建失败: {exc}") from exc


@router.delete("/mcp/servers/{name}")
# 方法作用：删除当前身份有权管理的 MCP Server。
# Args: name - Server 名称；scope - 可选精确作用域。
# Returns: 删除状态、名称和作用域。
async def delete_mcp_server(name: str, scope: str | None = None):
    """删除当前身份有权管理的 MCP Server，内置 YAML 服务不可删除。"""
    logger.debug("删除 MCP Server 入口", name=name, scope=scope or "")
    from src.api.auth import get_current_role, get_current_tenant_id, get_current_user_id
    from src.knowledge.governance import can_manage_knowledge_resource, normalize_knowledge_scope

    tenant_id = get_current_tenant_id()
    user_id = get_current_user_id()
    role = get_current_role()
    normalized_scope = None
    if scope:
        try:
            normalized_scope = normalize_knowledge_scope(scope).value
        except ValueError as exc:
            raise HTTPException(400, str(exc)) from exc
    try:
        conn = await _connect_scoped_mcp_db()
        try:
            row = await conn.fetchrow(
                "SELECT id, scope, tenant_id, owner_user_id, is_builtin FROM mcp_servers "
                "WHERE name=$1 AND ($2::text IS NULL OR scope=$2) AND "
                "(scope='system' OR (scope='tenant' AND tenant_id=$3) OR "
                "(scope='private' AND tenant_id=$3 AND owner_user_id=$4) OR $5='super_admin') "
                "ORDER BY CASE scope WHEN 'private' THEN 3 WHEN 'tenant' THEN 2 ELSE 1 END DESC "
                "LIMIT 1",
                name, normalized_scope, tenant_id, user_id, role,
            )
            if not row:
                raise HTTPException(404, f"MCP Server '{name}' 未找到")
            row_scope = str(row["scope"])
            row_tenant = int(row["tenant_id"] or 0)
            row_owner = int(row["owner_user_id"] or 0)
            if row["is_builtin"]:
                raise HTTPException(403, "内置 MCP Server 不可删除")
            if not can_manage_knowledge_resource(
                row_scope, role=role, current_tenant_id=tenant_id,
                resource_tenant_id=row_tenant, current_user_id=user_id,
                owner_user_id=row_owner,
            ):
                raise HTTPException(403, "无权删除该 MCP Server")
            await conn.execute("DELETE FROM mcp_servers WHERE id=$1", row["id"])
        finally:
            await conn.close()
        from src.mcp_client.client_manager import get_mcp_client_manager
        manager = get_mcp_client_manager()
        internal_name = manager.scoped_server_name(name, row_scope, row_tenant, row_owner)
        await manager.remove_server(internal_name)
        logger.info("删除 MCP Server 完成", name=name, scope=row_scope)
        return {"status": "ok", "name": name, "scope": row_scope}
    except HTTPException:
        raise
    except Exception as exc:
        logger.error("删除 MCP Server 失败", name=name, error=str(exc), exc_info=True)
        raise HTTPException(500, f"删除失败: {exc}") from exc


@router.post("/mcp/servers/{name}/test")
# 方法作用：测试当前身份可见 MCP Server 的连通性。
# Args: name - Server 名称；scope - 可选精确作用域。
# Returns: 测试结果和错误信息。
async def test_mcp_server(name: str, scope: str | None = None):
    """测试当前身份可见的 MCP Server 连通性。"""
    logger.debug("测试 MCP Server API 入口", name=name, scope=scope or "")
    try:
        from src.api.auth import get_current_tenant_id, get_current_user_id
        from src.mcp_client.client_manager import get_mcp_client_manager
        mgr = get_mcp_client_manager()
        tenant_id = get_current_tenant_id()
        user_id = get_current_user_id()
        await mgr.ensure_scoped_servers(tenant_id, user_id)
        candidates = []
        for internal_name in mgr.sessions:
            server_scope = mgr._server_scopes.get(internal_name, "system")  # noqa: SLF001
            if scope and server_scope != scope:
                continue
            if internal_name.endswith(f"_{name}") or internal_name == name:
                candidates.append(internal_name)
        if not candidates:
            raise HTTPException(404, f"MCP Server '{name}' 未找到")
        priority = {"private": 3, "tenant": 2, "system": 1}
        internal_name = max(
            candidates,
            key=lambda item: priority.get(mgr._server_scopes.get(item, "system"), 0),  # noqa: SLF001
        )
        ok = await mgr.test_connection(internal_name)
        result = {"name": name, "scope": mgr._server_scopes.get(internal_name, "system"), "ok": ok}  # noqa: SLF001
        logger.info("测试 MCP Server API 完成", name=name, ok=ok)
        return result
    except HTTPException:
        raise
    except Exception as e:
        logger.error("测试 MCP Server API 失败", name=name, error=str(e), exc_info=True)
        return {"name": name, "ok": False, "error": str(e)}


# ---- 查询历史 ----


@router.get("/history")
async def list_history(
    datasource: str | None = Query(default=None),
    search: str | None = Query(default=None),
    page: int = Query(default=1, ge=1),
    page_size: int = Query(default=50, ge=1, le=200),
):
    """分页列出查询历史，PG 持久化、重启不丢失。"""
    from src.memory.history_store import get_history_store
    return await get_history_store().list(
        datasource=datasource, search=search, page=page, page_size=page_size)


# ---- 会话管理 ----


@router.get("/sessions")
async def list_sessions(
    cursor: str | None = Query(default=None),
    limit: int = Query(default=20, ge=1, le=100),
):
    """游标分页列出历史会话，按最近活跃时间倒序。

    cursor 为上一页最后一条的 last_active_at ISO 字符串，首次传空。
    """
    from src.memory.session_store import get_session_store

    logger.debug("列出会话路由入口", cursor=cursor, limit=limit)
    items = await get_session_store().list(cursor=cursor, limit=limit + 1)
    has_more = len(items) > limit
    page_items = items[:limit]
    next_cursor = page_items[-1]["last_active_at"] if has_more and page_items else None
    logger.info(
        "列出会话路由完成", count=len(page_items), has_more=has_more,
        next_cursor=next_cursor,
    )
    return {"sessions": page_items, "next_cursor": next_cursor, "has_more": has_more}


@router.get("/sessions/{session_id}")
async def get_session(session_id: str):
    """获取会话详情，包含最近 20 轮对话 + 最新一轮的富数据。

    从 PG 读取会话元数据 + 从 LangGraph Checkpointer 读取对话内容。
    额外返回 latest_state 用于还原最后一轮的完整 UI（图表、数据表、分析结论等）。
    """
    from src.memory.session_store import get_session_store
    session = await get_session_store().get(session_id)
    if not session:
        raise HTTPException(404, f"会话 '{session_id}' 未找到")

    loaded_turns = await _load_session_turns(session_id, limit=21)
    has_more = len(loaded_turns) > 20
    turns = loaded_turns[-20:]
    # 提取最新一轮的富数据用于前端还原完整 UI
    latest_state = await _load_latest_state(session_id)
    if turns:
        # 持久化逐轮结果是权威数据，checkpoint 只补充缺失字段，避免贫化状态覆盖完整响应。
        latest_state = _merge_rich_result(
            turns[-1].get("final_result", {}) or {}, latest_state or {},
        )
    if turns and latest_state:
        # latest_state 只对应会话最后一轮，禁止向更早轮次扩散。
        turns[-1] = {
            **turns[-1],
            "sql": latest_state.get("sql", "") or turns[-1].get("sql", ""),
            "assistant_summary": (
                (latest_state.get("analysis", {}) or {}).get("summary", "")
                or turns[-1].get("assistant_summary", "")
            ),
            "final_result": latest_state,
        }
        logger.info(
            "会话最后一轮富数据合并完成",
            session_id=session_id[:20],
            turn_id=turns[-1].get("turn_id", 0),
            sql_statements=len(latest_state.get("sql_statements", []) or []),
            data_rows=len(latest_state.get("data", []) or []),
        )
    logger.info(
        "会话详情输出探针", session_id=session_id[:20],
        turns=len(turns), has_more=has_more,
        rich_turns=sum(1 for turn in turns if turn.get("final_result")),
        sql_turns=sum(1 for turn in turns if turn.get("sql")),
        data_rows=sum(
            len((turn.get("final_result", {}) or {}).get("data", []) or [])
            for turn in turns
        ),
    )
    return {
        "session": session, "turns": turns,
        "latest_state": latest_state, "has_more": has_more,
    }


@router.get("/sessions/{session_id}/turns")
async def list_session_turns(
    session_id: str,
    before: int | None = Query(default=None),
    limit: int = Query(default=20, ge=1, le=100),
):
    """瀑布流加载会话的对话轮次。

    before 为轮次序号，只返回该序号之前的更早轮次。
    不传 before 返回最新的 limit 条。
    """
    from src.memory.session_store import get_session_store
    if await get_session_store().get(session_id) is None:
        logger.warning("拒绝读取无权会话轮次", session_id=session_id[:20])
        raise HTTPException(404, f"会话 '{session_id}' 未找到")
    loaded_turns = await _load_session_turns(session_id, before=before, limit=limit + 1)
    has_more = len(loaded_turns) > limit
    turns = loaded_turns[-limit:]
    return {"turns": turns, "has_more": has_more}


# ---- 模型管理 ----

@router.get("/models")
async def list_models():
    from src.llm.model_registry import get_model_registry
    from src.config import get_settings
    items = []
    for m in get_model_registry().list_all():
        items.append({"id": m.model_id, "provider": m.provider, "name": m.display_name,
                      "context_window": m.capabilities.context_window,
                      "vision": m.capabilities.vision, "reasoning": m.capabilities.reasoning})
    return {"models": items, "default": get_settings().llm_model}


@router.post("/models/test")
async def test_model(req: dict):
    import time as _t
    from src.llm.client import get_provider
    try:
        p = get_provider(req.get("model_id", ""))
        s = _t.monotonic()
        await p.agenerate([{"role": "user", "content": "ping"}], max_tokens=1)
        return {"ok": True, "latency_ms": round((_t.monotonic() - s) * 1000)}
    except Exception as e:
        return {"ok": False, "error": str(e)}


@router.delete("/sessions/{session_id}")
async def delete_session(session_id: str):
    """删除会话记录（不删除 Checkpointer 中的 state）。"""
    from src.memory.session_store import get_session_store
    ok = await get_session_store().delete(session_id)
    if not ok:
        raise HTTPException(500, "删除失败")
    return {"status": "ok", "session_id": session_id}


# 方法作用：按字段优先级合并持久化富结果与 checkpoint 回退结果。
# Args: primary - 权威的逐轮持久化结果；fallback - checkpoint 或兼容路径结果。
# Returns: 合并后的结构化响应字典。
def _merge_rich_result(primary: dict | None, fallback: dict | None) -> dict:
    """以持久化响应为主，仅用回退响应补齐空的富数据字段。"""
    try:
        primary_value = primary if isinstance(primary, dict) else {}
        fallback_value = fallback if isinstance(fallback, dict) else {}
        logger.debug(
            "合并历史富结果入口",
            primary_keys=sorted(primary_value.keys()),
            fallback_keys=sorted(fallback_value.keys()),
        )
        merged = dict(fallback_value)
        merged.update(primary_value)
        rich_fields = (
            "sql", "sql_statements", "data", "row_count", "analysis", "chart",
            "sql_reasoning_content",
        )
        for field in rich_fields:
            value = primary_value.get(field)
            if value is None or value == "" or value == [] or value == {} or value == 0:
                if field in fallback_value and fallback_value[field] not in (None, "", [], {}, 0):
                    merged[field] = fallback_value[field]
        logger.info(
            "合并历史富结果完成",
            primary_rich=bool(primary_value),
            fallback_rich=bool(fallback_value),
            sql_statements=len(merged.get("sql_statements", []) or []),
            data_rows=len(merged.get("data", []) or []),
        )
        return merged
    except Exception as exc:
        logger.error("合并历史富结果失败", error=str(exc), exc_info=True)
        return dict(primary or fallback or {})


async def _load_checkpoint_tuple(session_id: str) -> object | None:
    """优先读取当前身份命名空间，并兼容迁移前的原始会话线程。

    Args:
        session_id: 对外会话 ID。

    Returns:
        最新 checkpoint tuple；两个线程均无状态时返回 None。
    """
    from src.api.auth import scope_thread_id
    from src.memory.checkpointer import get_checkpointer

    cp = await get_checkpointer()
    scoped_thread_id = scope_thread_id(session_id)
    logger.debug(
        "加载会话 checkpoint 入口", session_id=session_id[:20],
        scoped_thread_id=scoped_thread_id[-60:], checkpointer=type(cp).__name__,
    )
    try:
        scoped_config = {"configurable": {"thread_id": scoped_thread_id}}
        checkpoint_tuple = await cp.aget_tuple(scoped_config)
        if checkpoint_tuple:
            logger.info(
                "加载会话 checkpoint 完成", session_id=session_id[:20],
                source="scoped",
            )
            return checkpoint_tuple

        logger.info(
            "命名空间 checkpoint 无状态，回退旧会话线程",
            session_id=session_id[:20], legacy_thread_id=session_id[:20],
        )
        legacy_config = {"configurable": {"thread_id": session_id}}
        checkpoint_tuple = await cp.aget_tuple(legacy_config)
        logger.info(
            "加载会话 checkpoint 完成", session_id=session_id[:20],
            source="legacy" if checkpoint_tuple else "missing",
        )
        return checkpoint_tuple
    except Exception as exc:
        logger.error(
            "加载会话 checkpoint 失败", session_id=session_id[:20],
            error=str(exc), exc_info=True,
        )
        raise


async def _load_latest_state(session_id: str) -> dict | None:
    """从 Checkpointer 加载最新一轮的富数据（分析结论、图表、数据样本）。

    用于前端恢复历史会话时还原完整的分析结果 UI。
    """
    try:
        tup = await _load_checkpoint_tuple(session_id)
        if not tup:
            logger.info("最新状态 Checkpointer 无状态", session_id=session_id[:20])
            tup = None
        if tup:
            cv = tup.checkpoint.get("channel_values", {}) or {}
        else:
            cv = {}
        checkpoint_response = cv.get("final_response", {}) or {}
        logger.info(
            "历史最新状态字段探针",
            session_id=session_id[:20],
            channel_keys=sorted(str(key) for key in cv.keys()),
            has_final_response=bool(checkpoint_response),
            generated_sql=bool(cv.get("generated_sql")),
            final_sql=bool(checkpoint_response.get("sql"))
            if isinstance(checkpoint_response, dict) else False,
            final_sql_statements=len(checkpoint_response.get("sql_statements", []) or [])
            if isinstance(checkpoint_response, dict) else 0,
            final_data_rows=len(checkpoint_response.get("data", []) or [])
            if isinstance(checkpoint_response, dict) else 0,
        )
        # 多源查询的 generated_sql 可以为空，必须优先读取最终响应判断富数据。
        if isinstance(checkpoint_response, dict) and checkpoint_response:
            data_sample = checkpoint_response.get("data", cv.get("query_result_sample", [])) or []
            result = {
                "sql": checkpoint_response.get("sql", cv.get("generated_sql", "")) or "",
                "sql_statements": checkpoint_response.get("sql_statements", []) or [],
                "analysis": checkpoint_response.get("analysis", cv.get("analysis_result", {})) or {},
                "chart": checkpoint_response.get("chart", cv.get("chart_config", {})) or {},
                "data": data_sample if isinstance(data_sample, list) else [],
                "row_count": int(
                    checkpoint_response.get("row_count", cv.get("query_result_full_count", 0)) or 0
                ),
                "truncated": bool(
                    checkpoint_response.get("truncated", cv.get("query_result_truncated", False))
                ),
                "success": bool(checkpoint_response.get("success", True)),
                "error_message": checkpoint_response.get("error_message", "") or "",
                "sql_reasoning_content": checkpoint_response.get(
                    "sql_reasoning_content", cv.get("sql_reasoning_content", "")
                ) or "",
            }
            logger.info(
                "最新状态从最终响应恢复", session_id=session_id[:20],
                sql_statements=len(result["sql_statements"]),
                data_rows=len(result["data"]),
                has_analysis=bool(result["analysis"]),
            )
            return result

        # Checkpointer 不可用或没有最终响应时，回退到持久化逐轮查询历史。
        if not cv.get("generated_sql") and not cv.get("execution_error"):
            from src.memory.history_store import get_history_store
            history = await get_history_store().list_session(session_id, limit=1)
            if history:
                latest = history[-1]
                persisted_response = latest.get("final_result", {}) or {}
                if isinstance(persisted_response, dict) and persisted_response:
                    logger.info(
                        "最新状态从持久化结构化响应恢复",
                        session_id=session_id[:20],
                        data_rows=len(persisted_response.get("data", []) or []),
                    )
                    return persisted_response
                logger.info(
                    "最新状态从查询历史恢复", session_id=session_id[:20],
                    sql=bool(latest.get("sql")),
                )
                return {
                    "sql": latest.get("sql", "") or "",
                    "sql_statements": [],
                    "analysis": {}, "chart": {}, "data": [],
                    "row_count": int(latest.get("row_count", 0) or 0),
                    "truncated": False,
                    "success": bool(latest.get("success", True)),
                    "error_message": "" if latest.get("success", True) else "查询失败",
                    "sql_reasoning_content": "",
                }
        # 只提取前端需要的富数据字段
        data_sample = cv.get("query_result_sample", []) or []
        return {
            "sql": cv.get("generated_sql", "") or "",
            "sql_statements": [],
            "analysis": cv.get("analysis_result", {}) or {},
            "chart": cv.get("chart_config", {}) or {},
            "data": data_sample if isinstance(data_sample, list) else [],
            "row_count": int(cv.get("query_result_full_count", 0) or 0),
            "truncated": bool(cv.get("query_result_truncated", False)),
            "success": not cv.get("execution_error", ""),
            "error_message": cv.get("execution_error", "") or "",
            "sql_reasoning_content": cv.get("sql_reasoning_content", "") or "",
        }
    except Exception as e:
        logger.warning("最新状态加载失败", session_id=session_id[:20], error=str(e))
        return None


async def _load_session_turns(session_id: str, before: int | None = None, limit: int = 20) -> list[dict]:
    """从 LangGraph Checkpointer 加载会话的对话轮次。

    Args:
        session_id - 会话 ID（即 thread_id）
        before - 轮次序号游标，只返回此序号之前的轮次
        limit - 返回条数上限

    Returns: 对话轮次列表
    """
    try:
        logger.debug(
            "加载会话轮次边界", session_id=session_id[:20],
            before=before, limit=limit,
        )
        # 使用最新 checkpoint 恢复完整摘要，并用持久化历史补齐逐轮富数据。
        tup = await _load_checkpoint_tuple(session_id)
        channel_values: dict = {}
        if not tup:
            logger.info("Checkpointer 无状态", session_id=session_id[:20])
            messages = []
        else:
            channel_values = tup.checkpoint.get("channel_values", {}) or {}
            messages = channel_values.get("messages", []) or []
        checkpoint_history_probe = channel_values.get("conversation_history", []) or []
        logger.info(
            "历史轮次结构探针",
            session_id=session_id[:20],
            message_count=len(messages),
            history_count=len(checkpoint_history_probe),
            rich_history_count=sum(
                1 for item in checkpoint_history_probe
                if isinstance(item, dict) and bool(item.get("final_result"))
            ),
            checkpoint_final_response=bool(channel_values.get("final_response")),
        )
        logger.info("会话轮次加载", session_id=session_id[:20], msg_count=len(messages))

        turns: list[dict] = []
        checkpoint_history = channel_values.get("conversation_history", []) or []
        if checkpoint_history:
            for index, item in enumerate(checkpoint_history):
                value = item if isinstance(item, dict) else {
                    "turn_id": getattr(item, "turn_id", index + 1),
                    "user_query": getattr(item, "user_query", ""),
                    "generated_sql": getattr(item, "generated_sql", ""),
                    "analysis_summary": getattr(item, "analysis_summary", ""),
                    "timestamp": getattr(item, "timestamp", ""),
                    "final_result": getattr(item, "final_result", {}),
                }
                if not value.get("user_query"):
                    continue
                final_result = value.get("final_result", {}) or {}
                analysis = final_result.get("analysis", {}) if isinstance(final_result, dict) else {}
                turns.append({
                    "turn_id": value.get("turn_id", index + 1),
                    "user_query": value.get("user_query", "") or "",
                    "assistant_summary": (
                        value.get("analysis_summary", "")
                        or (analysis.get("summary", "") if isinstance(analysis, dict) else "")
                        or ""
                    ),
                    "sql": (
                        final_result.get("sql", "")
                        if isinstance(final_result, dict) else ""
                    ) or value.get("generated_sql", "") or "",
                    "timestamp": str(value.get("timestamp", "") or ""),
                    "final_result": final_result if isinstance(final_result, dict) else {},
                })
            if turns:
                logger.info(
                    "会话轮次从 checkpoint 历史恢复",
                    session_id=session_id[:20], turns=len(turns),
                )
        elif messages:
            i = 0
            while i < len(messages):
                msg = messages[i]
                if type(msg).__name__ == "HumanMessage":
                    user_query = msg.content if isinstance(msg.content, str) else str(msg.content)
                    sql = ""
                    summary = ""
                    j = i + 1
                    while j < len(messages) and j < i + 4:
                        nxt = messages[j]
                        if type(nxt).__name__ == "AIMessage":
                            content = nxt.content if isinstance(nxt.content, str) else str(nxt.content)
                            if content.startswith("SQL: "):
                                parts = content.split("\n结论: ", 1)
                                sql = parts[0][5:] if parts else ""
                                summary = parts[1] if len(parts) > 1 else content[5:]
                            else:
                                summary = content
                            i = j
                            break
                        j += 1
                    turns.append({
                        "turn_id": len(turns) + 1,
                        "user_query": user_query,
                        "assistant_summary": summary,
                        "sql": sql,
                        "timestamp": "",
                        "final_result": {},
                    })
                i += 1

        if not turns and channel_values.get("user_query"):
            analysis = channel_values.get("analysis_result", {}) or {}
            summary = analysis.get("summary", "") if isinstance(analysis, dict) else ""
            turns = [{
                "turn_id": 1,
                "user_query": channel_values.get("user_query", "") or "",
                "assistant_summary": (
                    summary or channel_values.get("execution_error", "") or "会话未完成"
                ),
                "sql": channel_values.get("generated_sql", "") or "",
                "timestamp": "",
                "final_result": {},
            }]
            logger.info("会话轮次从 checkpoint 输入恢复", session_id=session_id[:20])

        from src.memory.history_store import get_history_store
        # 无论 checkpoint 保存的是 dict 历史还是纯 messages，都读取逐轮 JSONB，
        # 这样旧 checkpoint 也能恢复 SQL、数据和图表；没有 checkpoint 时保留分页参数。
        history = await get_history_store().list_session(
            session_id,
            before=None if turns else before,
            limit=1000 if turns else limit,
        )
        persisted_turns = [{
                "turn_id": item.get("turn_id", index + 1),
                "user_query": item.get("query", "") or "",
                "assistant_summary": (
                    ((item.get("final_result", {}) or {}).get("analysis", {}) or {}).get("summary", "")
                    or (f"查询成功，返回 {item.get('row_count', 0) or 0} 行"
                        if item.get("success", True) else "查询失败")
                ),
                "sql": item.get("sql", "") or "",
                "timestamp": item.get("time", "") or "",
                "final_result": item.get("final_result", {}) or {},
            } for index, item in enumerate(history)]
        if not turns:
            turns = persisted_turns
            logger.info(
                "会话轮次从查询历史恢复", session_id=session_id[:20],
                turns=len(turns),
            )
        elif persisted_turns:
            persisted_by_id = {item["turn_id"]: item for item in persisted_turns}
            for turn in turns:
                persisted = persisted_by_id.get(turn["turn_id"])
                if not persisted:
                    continue
                if persisted.get("final_result"):
                    turn["final_result"] = _merge_rich_result(
                        persisted["final_result"], turn.get("final_result", {}) or {},
                    )
                turn["sql"] = persisted.get("sql") or turn.get("sql", "")
                turn["timestamp"] = persisted.get("timestamp") or turn.get("timestamp", "")
                if not turn.get("assistant_summary"):
                    turn["assistant_summary"] = persisted.get("assistant_summary", "")
            logger.info(
                "会话轮次结构化响应合并完成", session_id=session_id[:20],
                persisted_turns=len(persisted_turns),
                rich_turns=sum(1 for turn in turns if turn.get("final_result")),
            )

        # 旧记录没有富数据时构造最小响应，至少恢复完整文本与可用 SQL。
        for turn in turns:
            if turn.get("final_result"):
                continue
            turn["final_result"] = {
                "success": True,
                "sql": turn.get("sql", ""),
                "sql_statements": [],
                "data": [],
                "row_count": 0,
                "truncated": False,
                "analysis": {
                    "summary": turn.get("assistant_summary", ""),
                    "insights": [],
                    "recommended_chart_type": "table",
                },
                "chart": {"type": "table", "option": {}},
            }

        candidates = [turn for turn in turns if before is None or turn["turn_id"] < before]
        result = candidates[-limit:]

        logger.info(
            "会话轮次解析完成", session_id=session_id[:20],
            turns=len(result), total_candidates=len(candidates),
            rich_turns=sum(1 for turn in result if turn.get("final_result")),
        )
        return result
    except Exception as e:
        logger.warning("会话轮次加载失败", session_id=session_id[:20], error=str(e))
        return []


# ---- Skills 管理 ----


@router.get("/skills")
# 方法作用：列出当前身份可见的 Skill 资源。
# Args: skill_scope - 可选 system/tenant/private 过滤范围。
# Returns: Skill 列表和总数。
async def list_skills(skill_scope: str | None = None):
    """列出当前身份可见的系统、租户和个人 Skills。"""
    logger.debug("Skill 列表入口", skill_scope=skill_scope or "")
    try:
        from src.api.auth import get_current_tenant_id, get_current_user_id
        from src.knowledge.governance import normalize_knowledge_scope
        from src.skill_manager import get_skill_manager
        mgr = get_skill_manager()
        skills = []
        normalized_scope = None
        if skill_scope:
            try:
                normalized_scope = normalize_knowledge_scope(skill_scope).value
            except ValueError as exc:
                raise HTTPException(400, str(exc)) from exc
        tenant_id = get_current_tenant_id()
        user_id = get_current_user_id()
        for s in mgr.get_visible_skills(tenant_id, user_id):
            if normalized_scope and s.scope != normalized_scope:
                continue
            triggers = s.triggers or {}
            tools = s.tools or []
            deps = s.depends_on or {}
            skills.append({
                "name": s.name,
                "version": s.version,
                "enabled": s.enabled,
                "description": s.description or "",
                "triggers": triggers.get("keywords", []),
                "intents": triggers.get("intents", []),
                "tools": [t.get("name", "") for t in tools],
                "dependencies": deps.get("python_packages", []),
                "is_builtin": mgr.is_builtin(
                    s.name, tenant_id=tenant_id, user_id=user_id, scope=s.scope,
                ),
                "scope": s.scope,
                "tenant_id": s.tenant_id,
                "owner_user_id": s.owner_user_id,
            })
        result = {"skills": skills, "total": len(skills)}
        logger.info("Skill 列表完成", total=len(skills), tenant_id=tenant_id, user_id=user_id)
        return result
    except HTTPException:
        raise
    except Exception as e:
        logger.error("Skills 列表加载失败", error=str(e), exc_info=True)
        return {"skills": [], "total": 0}


@router.post("/skills/upload")
# 方法作用：把 Skill 清单上传到当前身份有权写入的受管目录。
# Args: files - 上传文件；skill_scope - 目标作用域，默认 private。
# Returns: 导入、跳过和失败明细。
async def upload_skills(
    files: list[UploadFile] = File(...),
    skill_scope: str = "private",
):
    """按当前认证身份批量上传 SKILL.md 到受管作用域目录。

    前端选择文件夹时，通过 webkitdirectory 传入整个文件夹的所有文件，
    后端递归过滤 SKILL.md（大小写不敏感），写入 skills/<name>/SKILL.md。
    """
    import re
    import yaml
    from pathlib import Path

    from src.api.auth import (
        get_current_role, get_current_tenant_id, get_current_user_id,
    )
    from src.config import get_settings
    from src.knowledge.governance import can_write_knowledge_scope, normalize_knowledge_scope
    from src.skill_manager import get_skill_manager

    logger.debug("Skill 上传入口", skill_scope=skill_scope, file_count=len(files))
    try:
        normalized_scope = normalize_knowledge_scope(skill_scope).value
    except ValueError as exc:
        raise HTTPException(400, str(exc)) from exc
    role = get_current_role()
    tenant_id = get_current_tenant_id()
    user_id = get_current_user_id()
    if not can_write_knowledge_scope(
        normalized_scope,
        role=role,
        user_id=user_id,
        multi_tenant=get_settings().multi_tenant,
    ):
        logger.warning(
            "Skill 上传权限拒绝",
            skill_scope=normalized_scope,
            role=role,
            tenant_id=tenant_id,
            user_id=user_id,
        )
        raise HTTPException(403, f"当前角色无权写入 {normalized_scope} Skill")
    mgr = get_skill_manager()
    skills_dir = str(mgr.get_upload_dir(
        normalized_scope, tenant_id=tenant_id, user_id=user_id,
    ))
    os.makedirs(skills_dir, exist_ok=True)
    imported: list[dict] = []
    skipped: list[str] = []
    errors: list[dict] = []

    for f in files:
        if not f.filename:
            continue
        # 匹配路径中各级目录下的 SKILL.md（大小写不敏感）
        basename = os.path.basename(f.filename).lower()
        if basename != "skill.md":
            skipped.append(f.filename)
            continue
        try:
            content = (await f.read()).decode("utf-8")
        except Exception as e:
            errors.append({"file": f.filename, "error": str(e)})
            continue

        # 解析 YAML frontmatter 提取 name
        fm_match = re.match(r"^---\s*\n(.*?)\n---", content, re.DOTALL)
        skill_name = None
        if fm_match:
            try:
                fm = yaml.safe_load(fm_match.group(1))
                skill_name = fm.get("name", "")
            except Exception:
                pass
        if not skill_name:
            # 回退：从父目录名推断
            parent = os.path.basename(os.path.dirname(f.filename))
            if parent and parent != ".":
                skill_name = parent
            else:
                errors.append({"file": f.filename, "error": "无法从 frontmatter 或路径提取 skill name"})
                continue

        # 写文件
        dest_dir = os.path.realpath(os.path.join(skills_dir, os.path.basename(skill_name)))
        if os.path.commonpath([dest_dir, os.path.realpath(skills_dir)]) != os.path.realpath(skills_dir):
            raise HTTPException(403, "禁止访问")
        os.makedirs(dest_dir, exist_ok=True)
        dest = os.path.join(dest_dir, "SKILL.md")
        with open(dest, "w", encoding="utf-8") as fp:
            fp.write(content)
        imported.append({
            "name": skill_name, "file": f.filename, "scope": normalized_scope,
        })

        # 递归检测：同级目录下有 tools.py 或 templates/ 也一并说明
        # （前端传了整个文件夹时这些文件也会在 files 中，但会被 basename 检查跳过）

    # 将已导入的 skill 注入缓存（不扫描整个文件夹，增量更新）
    for item in imported:
        try:
            from src.skill_manager import get_skill_manager
            mgr2 = get_skill_manager()
            md_path = os.path.join(skills_dir, item["name"], "SKILL.md")
            if os.path.isfile(md_path):
                skill_obj = mgr2._parse_skill_manifest(  # noqa: SLF001
                    Path(md_path),
                    scope=normalized_scope,
                    tenant_id=tenant_id if normalized_scope != "system" else 0,
                    owner_user_id=user_id if normalized_scope == "private" else 0,
                )
                mgr2.add_skill(skill_obj)
        except Exception as exc:
            logger.error(
                "Skill 上传缓存注入失败", name=item["name"], error=str(exc), exc_info=True,
            )
            errors.append({"file": item["file"], "error": str(exc)})

    result = {
        "imported": imported,
        "skipped": skipped,
        "errors": errors,
        "total": len(imported),
    }
    logger.info(
        "Skill 上传完成",
        skill_scope=normalized_scope,
        tenant_id=tenant_id,
        user_id=user_id,
        imported=len(imported),
        errors=len(errors),
    )
    return result


# ---- Knowledge / 知识库 ----


@router.post("/assets/profile")
async def profile_structured_asset(file: UploadFile = File(...)):
    """解析 CSV/Excel/Parquet 并返回列级 profile，不把原文件写入知识库。"""
    from src.config import get_settings
    from src.knowledge.structured_assets import StructuredAssetAdapter, StructuredAssetError

    logger.debug("结构化资产 profile API 入口", file_name=file.filename or "")
    if not file.filename:
        raise HTTPException(400, "文件名不能为空")
    settings = get_settings()
    max_bytes = settings.max_upload_bytes
    try:
        content = await file.read(max_bytes + 1)
        profile = StructuredAssetAdapter(max_bytes=max_bytes).inspect_bytes(file.filename, content)
        result = profile.to_dict()
        logger.info("结构化资产 profile API 完成", file_name=file.filename, rows=profile.row_count)
        return result
    except StructuredAssetError as exc:
        logger.warning("结构化资产 profile API 拒绝", file_name=file.filename, error=str(exc))
        raise HTTPException(400, str(exc)) from exc
    except Exception as exc:
        logger.error("结构化资产 profile API 失败", file_name=file.filename, error=str(exc), exc_info=True)
        raise HTTPException(500, "结构化资产解析失败") from exc


@router.post("/assets/query")
async def query_structured_asset(
    file: UploadFile = File(...),
    sql: str = Query(..., min_length=1),
    sheet_name: str | None = Query(default=None),
):
    """对上传的 CSV/Excel/Parquet 执行受控只读 SQL。"""
    from src.config import get_settings
    from src.knowledge.structured_query import StructuredQueryEngine, StructuredQueryError

    logger.debug("结构化资产查询 API 入口", file_name=file.filename or "", sql_preview=sql[:120])
    if not file.filename:
        raise HTTPException(400, "文件名不能为空")
    settings = get_settings()
    max_bytes = settings.max_upload_bytes
    try:
        content = await file.read(max_bytes + 1)
        result = await StructuredQueryEngine(
            max_rows=settings.max_result_rows,
            max_bytes=max_bytes,
            max_scan_rows=getattr(settings, "max_scan_rows", 1_000_000),
        ).execute(file.filename, content, sql, sheet_name=sheet_name)
        logger.info("结构化资产查询 API 完成", file_name=file.filename, rows=result.row_count)
        return result.to_dict()
    except StructuredQueryError as exc:
        logger.warning("结构化资产查询 API 拒绝", file_name=file.filename, error=str(exc))
        raise HTTPException(400, str(exc)) from exc
    except Exception as exc:
        logger.error("结构化资产查询 API 失败", file_name=file.filename, error=str(exc), exc_info=True)
        raise HTTPException(500, "结构化资产查询失败") from exc


@router.post("/analysis/forecast")
async def forecast_asset(payload: dict = Body(...)):
    """对已提供的时间序列行执行带回测和区间的确定性预测。"""
    from src.tools.forecasting import ForecastingError, forecast_rows

    logger.debug("预测 API 入口", payload_keys=sorted(payload.keys()))
    try:
        rows = payload.get("rows", [])
        result = forecast_rows(
            rows=rows,
            time_col=str(payload.get("time_col", "")),
            value_col=str(payload.get("value_col", "")),
            horizon=int(payload.get("horizon", 3)),
        )
        logger.info("预测 API 完成", model=result.model, horizon=len(result.predictions))
        return result.to_dict()
    except ForecastingError as exc:
        logger.warning("预测 API 输入拒绝", error=str(exc))
        raise HTTPException(400, str(exc)) from exc
    except Exception as exc:
        logger.error("预测 API 失败", error=str(exc), exc_info=True)
        raise HTTPException(500, "预测执行失败") from exc


@router.get("/knowledge")
async def list_knowledge(
    category: str | None = Query(default=None),
    search: str | None = Query(default=None),
    knowledge_scope: str | None = None,
    page: int = Query(default=1, ge=1),
    page_size: int = Query(default=20, ge=10, le=100),
):
    """按当前身份列出系统、租户和个人知识条目。"""
    logger.debug(
        "知识库列表入口",
        category=category or "",
        search=search or "",
        knowledge_scope=knowledge_scope or "",
        page=page,
        page_size=page_size,
    )
    try:
        from src.memory.vector_store import get_vector_store
        from src.knowledge.governance import normalize_knowledge_scope
        from src.knowledge.retrieval import build_accessible_knowledge_filters

        store = await get_vector_store()
        filter_groups = build_accessible_knowledge_filters(category=category or "")
        if knowledge_scope:
            normalized_scope = normalize_knowledge_scope(knowledge_scope).value
            filter_groups = [
                filters for filters in filter_groups
                if filters.get("visibility") == normalized_scope
            ]
        results_by_id = {}
        for filters in filter_groups:
            scoped_results = await store.get_by_filter(filters, limit=10000)
            logger.info(
                "知识库分范围查询完成",
                visibility=filters.get("visibility", ""),
                count=len(scoped_results),
            )
            for result in scoped_results:
                results_by_id[result.id] = result
        results = list(results_by_id.values())
        logger.info(
            "知识库查询完成",
            total_ids=len(results),
            scope_count=len(filter_groups),
        )
        entries = []
        from src.api.auth import get_current_role, get_current_tenant_id, get_current_user_id
        from src.knowledge.governance import can_manage_knowledge_resource

        for result in results:
            meta = dict(result.metadata or {})
            is_user = meta.get("source") == "user_upload"
            visibility = str(meta.get("visibility", "system" if not is_user else "tenant"))
            resource_tenant_id = meta.get("tenant_id")
            if resource_tenant_id is None and visibility != "system":
                resource_tenant_id = get_current_tenant_id()
            can_delete = is_user and can_manage_knowledge_resource(
                visibility,
                role=get_current_role(),
                current_tenant_id=get_current_tenant_id(),
                resource_tenant_id=resource_tenant_id,
                current_user_id=get_current_user_id(),
                owner_user_id=meta.get("owner_user_id"),
            )
            entries.append({
                "id": result.id,
                "content": result.content,
                "category": meta.get("category", ""),
                "datasource": meta.get("datasource", ""),
                "table_name": meta.get("table_name", ""),
                "source": meta.get("source", "unknown"),
                "source_file": meta.get("source_file", ""),
                "scope": visibility,
                "tag_ids": json.loads(meta.get("tag_ids_json", "[]") or "[]"),
                "tags": [tag for tag in str(meta.get("tags", "")).split(",") if tag],
                "is_builtin": meta.get("visibility") == "system" or not is_user,
                "tenant_id": resource_tenant_id,
                "owner_user_id": meta.get("owner_user_id"),
                "can_delete": can_delete,
            })
        if search:
            q = search.lower()
            entries = [e for e in entries if q in e["content"].lower()
                       or q in e.get("table_name", "").lower()]
        total = len(entries)
        start = (page - 1) * page_size
        paged = entries[start:start + page_size]
        result = {"entries": paged, "total": total, "page": page, "page_size": page_size}
        logger.info("知识库列表完成", total=total, returned=len(paged))
        return result
    except ValueError as exc:
        logger.warning("知识库列表范围无效", error=str(exc))
        raise HTTPException(400, str(exc)) from exc
    except Exception as e:
        logger.error("知识库加载失败", error=str(e), exc_info=True)
        return {"entries": [], "total": 0, "page": page, "page_size": page_size}


@router.post("/knowledge/docs/upload")
async def upload_knowledge_docs(
    files: list[UploadFile] = File(...),
    strategy: str = Query(default="auto"),
    chunk_size: int = Query(default=800, ge=200, le=4000),
    chunk_overlap: int = Query(default=100, ge=0, le=500),
    category: str = Query(default=""),
    knowledge_scope: str = "private",
    tag_ids: str = "",
    datasource: str = "",
):
    """按范围和标签批量上传文档并异步索引。

    system 仅超级管理员，tenant 仅租户管理员或超级管理员，private 归当前用户。
    """
    import asyncio as _asyncio
    from src.config import get_settings
    from src.api.auth import (
        get_current_role, get_current_tenant_id, get_current_user_id,
    )
    from src.knowledge.governance import can_write_knowledge_scope, normalize_knowledge_scope

    settings = get_settings()
    max_upload_bytes = settings.max_upload_bytes
    try:
        normalized_scope = normalize_knowledge_scope(knowledge_scope).value
    except ValueError as exc:
        logger.warning("知识文件上传范围无效", knowledge_scope=knowledge_scope)
        raise HTTPException(400, str(exc)) from exc
    role = get_current_role()
    user_id = get_current_user_id()
    if not can_write_knowledge_scope(
        normalized_scope,
        role=role,
        user_id=user_id,
        multi_tenant=getattr(settings, "multi_tenant", False),
    ):
        logger.warning(
            "知识文件上传权限拒绝",
            knowledge_scope=normalized_scope,
            role=role,
            user_id=user_id,
        )
        raise HTTPException(403, f"当前角色无权写入 {normalized_scope} 知识")
    logger.info(
        "知识文件上传授权完成",
        knowledge_scope=normalized_scope,
        role=role,
        tenant_id=get_current_tenant_id(),
        user_id=user_id,
    )

    try:
        requested_tag_ids = list(dict.fromkeys(
            int(value.strip()) for value in tag_ids.split(",") if value.strip()
        ))
    except ValueError as exc:
        logger.warning("知识文件上传标签格式无效", tag_ids=tag_ids)
        raise HTTPException(400, "tag_ids 必须是逗号分隔的整数") from exc
    if len(requested_tag_ids) > 20:
        raise HTTPException(400, "单个文档最多选择 20 个标签")
    resolved_tags = []
    if requested_tag_ids:
        from src.knowledge.tag_store import get_knowledge_tag_store

        resolved_tags = await get_knowledge_tag_store().get_visible_by_ids(
            requested_tag_ids,
            tenant_id=get_current_tenant_id(),
            user_id=user_id,
        )
        if {tag.id for tag in resolved_tags} != set(requested_tag_ids):
            logger.warning(
                "知识文件上传标签不可见",
                requested=requested_tag_ids,
                resolved=[tag.id for tag in resolved_tags],
            )
            raise HTTPException(400, "包含不存在、已停用或无权使用的标签")
        if normalized_scope != "private" and any(tag.scope != "global" for tag in resolved_tags):
            logger.warning(
                "公共知识上传标签范围拒绝",
                knowledge_scope=normalized_scope,
                private_tag_ids=[tag.id for tag in resolved_tags if tag.scope != "global"],
            )
            raise HTTPException(400, "系统或租户公共知识只能使用全局标签")
    resolved_tag_names = [tag.name for tag in resolved_tags]
    logger.debug(
        "知识文件上传入口",
        file_count=len(files),
        max_upload_bytes=max_upload_bytes,
        knowledge_scope=normalized_scope,
        tag_count=len(resolved_tags),
        datasource=datasource,
    )

    supported_exts = {".md", ".txt", ".pdf", ".docx", ".doc", ".markdown", ".csv"}
    from src.knowledge.doc_parser import ChunkConfig, ChunkStrategy
    from src.knowledge.upload_manager import get_upload_manager
    mgr = get_upload_manager()
    tasks_result: list[dict] = []
    errors: list[dict] = []

    try:
        strat = ChunkStrategy(strategy)
    except ValueError:
        strat = ChunkStrategy.AUTO
    config = ChunkConfig(strategy=strat, chunk_size=chunk_size, chunk_overlap=chunk_overlap)

    for f in files:
        if not f.filename:
            continue
        ext = os.path.splitext(f.filename)[1].lower()
        if ext not in supported_exts:
            errors.append({"file": f.filename, "error": f"不支持的文件格式: {ext}"})
            continue
        try:
            content = await f.read(max_upload_bytes + 1)
            if len(content) > max_upload_bytes:
                logger.warning("知识文件超过大小限制", filename=f.filename, size=len(content))
                raise HTTPException(413, f"文件 '{f.filename}' 超过大小限制")
            # 保存原始文件到 PostgreSQL
            from src.knowledge.file_store import get_file_store
            file_id = await get_file_store().save(
                f.filename,
                content,
                knowledge_scope=normalized_scope,
                datasource=datasource,
                tag_ids=requested_tag_ids,
            )

            task = mgr.create(
                f.filename,
                knowledge_scope=normalized_scope,
                tag_ids=requested_tag_ids,
                tag_names=resolved_tag_names,
                datasource=datasource,
            )
            tasks_result.append({"task_id": task.id, "file_name": f.filename, "file_id": file_id})
            _asyncio.create_task(mgr.process(task, content, config, category))
        except HTTPException:
            raise
        except Exception as e:
            errors.append({"file": f.filename, "error": str(e)})

    result = {
        "tasks": tasks_result, "errors": errors, "total": len(tasks_result),
        "config": config.to_dict(),
        "message": "已接收文件，后台处理中。轮询 GET /api/v1/knowledge/upload/status 获取进度",
    }
    logger.info(
        "知识文件上传入口完成",
        accepted=len(tasks_result),
        errors=len(errors),
        knowledge_scope=normalized_scope,
    )
    return result


@router.get("/knowledge/upload/status")
async def upload_status(task_id: str = Query(default="")):
    """查询上传任务进度（单任务或全部最近任务）。"""
    from src.knowledge.upload_manager import get_upload_manager
    mgr = get_upload_manager()
    if task_id:
        t = mgr.get(task_id)
        if t is None:
            raise HTTPException(404, f"任务 '{task_id}' 未找到")
        return {"task": t.to_dict()}
    return {"tasks": mgr.list_recent()}


@router.get("/knowledge/tags")
async def search_knowledge_tags(
    q: str = Query(default="", max_length=128),
    include_inactive: bool = Query(default=False),
    limit: int = Query(default=30, ge=1, le=100),
):
    """搜索当前用户可见的全局标签和个人标签。"""
    from src.api.auth import (
        get_current_role, get_current_tenant_id, get_current_user_id,
    )
    from src.knowledge.governance import is_super_admin
    from src.knowledge.tag_store import get_knowledge_tag_store

    role = get_current_role()
    effective_include_inactive = include_inactive and is_super_admin(role)
    logger.debug(
        "知识标签搜索 API 入口",
        query=q,
        include_inactive=effective_include_inactive,
        limit=limit,
        role=role,
    )
    try:
        tags = await get_knowledge_tag_store().search(
            q,
            tenant_id=get_current_tenant_id(),
            user_id=get_current_user_id(),
            limit=limit,
            include_inactive=effective_include_inactive,
            include_all_private=is_super_admin(role),
        )
        result = {"tags": [tag.model_dump() for tag in tags], "total": len(tags)}
        logger.info("知识标签搜索 API 完成", count=len(tags), role=role)
        return result
    except Exception as exc:
        logger.error("知识标签搜索 API 失败", error=str(exc), exc_info=True)
        raise HTTPException(500, "知识标签搜索失败") from exc


@router.post("/knowledge/tags", status_code=201)
async def create_knowledge_tag(request: KnowledgeTagCreateRequest):
    """为当前用户创建默认仅本人可见的自定义标签。"""
    from src.api.auth import get_current_tenant_id, get_current_user_id
    from src.knowledge.tag_store import get_knowledge_tag_store

    tenant_id = get_current_tenant_id()
    user_id = get_current_user_id()
    logger.debug("创建个人知识标签 API 入口", name=request.name, tenant_id=tenant_id, user_id=user_id)
    try:
        tag = await get_knowledge_tag_store().create_personal(
            request.name,
            tenant_id=tenant_id,
            user_id=user_id,
            description=request.description,
            aliases=request.aliases,
        )
        result = tag.model_dump()
        logger.info("创建个人知识标签 API 完成", tag_id=result["id"], user_id=user_id)
        return result
    except ValueError as exc:
        logger.warning("创建个人知识标签 API 输入拒绝", error=str(exc), user_id=user_id)
        raise HTTPException(400, str(exc)) from exc
    except Exception as exc:
        logger.error("创建个人知识标签 API 失败", error=str(exc), exc_info=True)
        raise HTTPException(500, "创建个人标签失败") from exc


@router.post("/knowledge/tags/global", status_code=201)
async def create_global_knowledge_tag(request: KnowledgeTagCreateRequest):
    """由超级管理员创建或重新启用平台全局标签。"""
    from src.api.auth import get_current_role, require_super_admin
    from src.knowledge.tag_store import get_knowledge_tag_store

    logger.debug("创建全局知识标签 API 入口", name=request.name, role=get_current_role())
    require_super_admin()
    try:
        tag = await get_knowledge_tag_store().create_global(
            request.name,
            actor_role=get_current_role(),
            tag_group=request.tag_group,
            description=request.description,
            aliases=request.aliases,
        )
        result = tag.model_dump()
        logger.info("创建全局知识标签 API 完成", tag_id=result["id"])
        return result
    except ValueError as exc:
        logger.warning("创建全局知识标签 API 输入拒绝", error=str(exc))
        raise HTTPException(400, str(exc)) from exc
    except Exception as exc:
        logger.error("创建全局知识标签 API 失败", error=str(exc), exc_info=True)
        raise HTTPException(500, "创建全局标签失败") from exc


@router.post("/knowledge/tags/{tag_id}/promote")
async def promote_knowledge_tag(tag_id: int):
    """由超级管理员把个人标签提升为平台全局标签。"""
    from src.api.auth import get_current_role, require_super_admin
    from src.knowledge.tag_store import get_knowledge_tag_store

    logger.debug("提升知识标签 API 入口", tag_id=tag_id, role=get_current_role())
    require_super_admin()
    try:
        tag = await get_knowledge_tag_store().promote_to_global(
            tag_id,
            actor_role=get_current_role(),
        )
        result = tag.model_dump()
        logger.info("提升知识标签 API 完成", source_tag_id=tag_id, global_tag_id=result["id"])
        return result
    except ValueError as exc:
        logger.warning("提升知识标签 API 未命中", tag_id=tag_id, error=str(exc))
        raise HTTPException(404, str(exc)) from exc
    except Exception as exc:
        logger.error("提升知识标签 API 失败", tag_id=tag_id, error=str(exc), exc_info=True)
        raise HTTPException(500, "提升标签失败") from exc


@router.patch("/knowledge/tags/{tag_id}")
async def update_knowledge_tag_status(tag_id: int, request: KnowledgeTagStatusRequest):
    """启用或停用全局标签或当前用户自己的个人标签。"""
    from src.api.auth import (
        get_current_role, get_current_tenant_id, get_current_user_id,
    )
    from src.knowledge.tag_store import get_knowledge_tag_store

    logger.debug(
        "更新知识标签状态 API 入口",
        tag_id=tag_id,
        is_active=request.is_active,
        role=get_current_role(),
    )
    try:
        updated = await get_knowledge_tag_store().set_active(
            tag_id,
            request.is_active,
            actor_role=get_current_role(),
            tenant_id=get_current_tenant_id(),
            user_id=get_current_user_id(),
        )
        if not updated:
            raise HTTPException(404, f"标签 '{tag_id}' 不存在或无权管理")
        logger.info("更新知识标签状态 API 完成", tag_id=tag_id, is_active=request.is_active)
        return {"status": "ok", "id": tag_id, "is_active": request.is_active}
    except HTTPException:
        raise
    except Exception as exc:
        logger.error("更新知识标签状态 API 失败", tag_id=tag_id, error=str(exc), exc_info=True)
        raise HTTPException(500, "更新标签状态失败") from exc


@router.get("/knowledge/test-search")
async def test_knowledge_search(q: str = Query(default=""), datasource: str = Query(default="")):
    """测试知识库检索效果——验证上传的文档是否能被检索到。

    有 q 时做语义搜索，返回匹配条目和相关性分数。
    无 q 时返回全部条目列表。
    """
    try:
        from src.memory.vector_store import get_vector_store
        from src.knowledge.retrieval import build_knowledge_filters, search_knowledge
        store = await get_vector_store()
        from src.api.auth import get_current_tenant_id
        if q:
            evidence = await search_knowledge(store, q, datasource=datasource, top_k=10)
            items = [{"rank": i + 1, "id": e.source_id, "content": e.content[:200],
                      "relevance": e.scores.get("relevance", 0.0), "citation": e.locator}
                     for i, e in enumerate(evidence)]
            return {"query": q, "results": items, "total": len(items)}
        else:
            filters = build_knowledge_filters(datasource=datasource)
            results = await store.get_by_filter(filters, limit=1000)
            return {"total": len(results), "ids": [r.id for r in results]}
    except Exception as e:
        raise HTTPException(500, f"知识库检索测试失败: {e}")


@router.get("/knowledge/docs")
async def list_knowledge_docs():
    """列出当前身份可见文档并附带服务端删除权限。"""
    from src.api.auth import get_current_role, get_current_tenant_id, get_current_user_id
    from src.knowledge.governance import can_manage_knowledge_resource
    from src.knowledge.file_store import get_file_store

    logger.debug("知识文档列表 API 入口", role=get_current_role())
    docs = await get_file_store().list_files()
    for doc in docs:
        scope = str(doc.get("scope", "system" if doc.get("is_builtin") else "private"))
        doc["can_delete"] = not doc.get("is_builtin", False) and can_manage_knowledge_resource(
            scope,
            role=get_current_role(),
            current_tenant_id=get_current_tenant_id(),
            resource_tenant_id=doc.get("tenant_id", get_current_tenant_id()),
            current_user_id=get_current_user_id(),
            owner_user_id=doc.get("owner_user_id"),
        )
    result = {"docs": docs, "total": len(docs)}
    logger.info("知识文档列表 API 完成", total=len(docs), role=get_current_role())
    return result


@router.get("/knowledge/docs/{doc_name}/content")
async def get_doc_content(doc_name: str, knowledge_scope: str = ""):
    """获取已索引文档的内容（从 PG 读取，回退磁盘）。"""
    from src.knowledge.file_store import get_file_store
    store = get_file_store()
    doc = await store.get_by_name(doc_name, knowledge_scope=knowledge_scope)
    if doc:
        raw = doc["file_data"]
        size = doc["size"]
    else:
        if knowledge_scope and knowledge_scope != "system":
            raise HTTPException(404, f"文档 '{doc_name}' 未找到")
        # 磁盘回退仅用于系统内置目录。
        base_dir = os.path.realpath(os.path.join(os.path.dirname(__file__), "..", "..", "docs", "metrics"))
        doc_path = os.path.realpath(os.path.join(base_dir, os.path.basename(doc_name)))
        if not doc_path.startswith(base_dir):
            raise HTTPException(403, "禁止访问")
        if not os.path.isfile(doc_path):
            raise HTTPException(404, f"文档 '{doc_name}' 未找到")
        with open(doc_path, "rb") as fp:
            raw = fp.read()
        size = os.path.getsize(doc_path)

    ext = os.path.splitext(doc_name)[1].lower()
    result = {"name": doc_name, "size": size, "ext": ext, "type": "text", "content": "", "raw_url": ""}
    if ext == ".pdf":
        result["type"] = "pdf"
        result["raw_url"] = (
            f"/api/v1/knowledge/docs/{doc_name}/raw"
            f"?knowledge_scope={knowledge_scope or 'system'}"
        )
        from src.knowledge.doc_parser import extract_text
        result["content"] = extract_text(doc_name, raw)
    elif ext in (".docx", ".doc"):
        result["type"] = "word"
        result["content"] = _docx_to_html(raw)
    else:
        result["content"] = raw.decode("utf-8", errors="replace")
    return result


@router.get("/knowledge/docs/{doc_name}/raw")
async def get_doc_raw(doc_name: str, knowledge_scope: str = ""):
    """返回原始文件（从 PG 读取，回退磁盘，用于 PDF iframe 渲染）。"""
    from fastapi.responses import Response
    from src.knowledge.file_store import get_file_store
    doc = await get_file_store().get_by_name(doc_name, knowledge_scope=knowledge_scope)
    if doc:
        return Response(content=bytes(doc["file_data"]), media_type=doc["content_type"])
    if knowledge_scope and knowledge_scope != "system":
        raise HTTPException(404, f"文档 '{doc_name}' 未找到")
    # 磁盘回退仅用于系统内置目录。
    from fastapi.responses import FileResponse
    base_dir = os.path.realpath(os.path.join(os.path.dirname(__file__), "..", "..", "docs", "metrics"))
    doc_path = os.path.realpath(os.path.join(base_dir, os.path.basename(doc_name)))
    if not doc_path.startswith(base_dir):
        raise HTTPException(403, "禁止访问")
    if not os.path.isfile(doc_path):
        raise HTTPException(404, f"文档 '{doc_name}' 未找到")
    ext = os.path.splitext(doc_name)[1].lower()
    media_map = {".pdf": "application/pdf", ".txt": "text/plain", ".md": "text/markdown"}
    return FileResponse(doc_path, media_type=media_map.get(ext, "application/octet-stream"))


def _docx_to_html(raw: bytes) -> str:
    """将 Word 文档转换为转义后的简单 HTML。

    Args:
        raw: Word 文档二进制内容。

    Returns:
        可安全嵌入前端的 HTML 文本。
    """
    logger.debug("Word HTML 转换入口", size=len(raw))
    try:
        from docx import Document
        doc = Document(io.BytesIO(raw))
        parts = ['<div style="font-family: sans-serif; line-height: 1.8;">']
        for para in doc.paragraphs:
            text = html.escape(para.text.strip(), quote=True)
            if not text:
                parts.append('<br/>')
                continue
            if para.style and para.style.name and para.style.name.startswith("Heading"):
                level = para.style.name.replace("Heading", "").strip()
                try:
                    lv = max(1, min(6, int(level)))
                except ValueError:
                    lv = 2
                parts.append(f'<h{lv} style="margin:12px 0 6px;">{text}</h{lv}>')
            else:
                parts.append(f'<p style="margin:4px 0;">{text}</p>')
        # 处理表格
        for table in doc.tables:
            parts.append('<table border="1" cellpadding="6" cellspacing="0" style="border-collapse:collapse; width:100%; margin:8px 0;">')
            for row in table.rows:
                parts.append('<tr>')
                for cell in row.cells:
                    cell_text = html.escape(cell.text, quote=True)
                    parts.append(f'<td style="padding:4px 8px;">{cell_text}</td>')
                parts.append('</tr>')
            parts.append('</table>')
        parts.append('</div>')
        result = "\n".join(parts)
        logger.info("Word HTML 转换完成", size=len(raw), output_size=len(result))
        return result
    except Exception as exc:
        logger.error("Word HTML 转换失败", error=str(exc), exc_info=True)
        return "<p>无法渲染 Word 文档内容</p>"


@router.delete("/knowledge/{entry_id}")
async def delete_knowledge_entry(entry_id: str):
    """按知识范围和管理员角色删除指定用户上传条目。"""
    logger.debug("删除知识条目 API 入口", entry_id=entry_id)
    try:
        from src.memory.vector_store import get_vector_store
        store = await get_vector_store()
        entry = await store.get_by_id(entry_id)
        if entry is None:
            logger.warning("删除知识条目未命中", entry_id=entry_id)
            raise HTTPException(404, f"知识条目 '{entry_id}' 未找到")
        meta = dict(entry.metadata or {})
        if meta.get("source") != "user_upload":
            logger.warning("删除知识条目拒绝", entry_id=entry_id, reason="非用户上传")
            raise HTTPException(403, "目录扫描或系统内置知识条目不可通过 API 删除")
        from src.api.auth import (
            get_current_role, get_current_tenant_id, get_current_user_id,
        )
        from src.knowledge.governance import can_manage_knowledge_resource

        visibility = str(meta.get("visibility") or "private")
        current_tenant_id = get_current_tenant_id()
        resource_tenant_id = meta.get("tenant_id")
        if resource_tenant_id is None and visibility != "system":
            resource_tenant_id = current_tenant_id
        owner_user_id = meta.get("owner_user_id")
        allowed = can_manage_knowledge_resource(
            visibility,
            role=get_current_role(),
            current_tenant_id=current_tenant_id,
            resource_tenant_id=resource_tenant_id,
            current_user_id=get_current_user_id(),
            owner_user_id=owner_user_id,
        )
        if not allowed:
            logger.warning(
                "删除知识条目权限拒绝",
                entry_id=entry_id,
                visibility=visibility,
                role=get_current_role(),
            )
            raise HTTPException(403, "无权删除该知识条目")
        await store.delete_by_ids([entry_id])
        result = {"status": "ok", "id": entry_id}
        logger.info("删除知识条目 API 完成", entry_id=entry_id, visibility=visibility)
        return result
    except HTTPException:
        raise
    except Exception as e:
        logger.error("删除知识条目 API 失败", entry_id=entry_id, error=str(e), exc_info=True)
        raise HTTPException(500, str(e))


@router.delete("/knowledge/docs/{doc_name}")
async def delete_knowledge_doc(doc_name: str, knowledge_scope: str = ""):
    """按范围删除当前身份可管理的原文件及全部关联向量。"""
    logger.debug(
        "删除知识文档 API 入口",
        doc_name=doc_name,
        knowledge_scope=knowledge_scope,
    )
    from src.memory.vector_store import get_vector_store
    from src.api.auth import get_current_role, get_current_tenant_id, get_current_user_id
    from src.knowledge.governance import can_manage_knowledge_resource, normalize_knowledge_scope
    from src.knowledge.retrieval import build_accessible_knowledge_filters

    if knowledge_scope:
        try:
            normalized_scope = normalize_knowledge_scope(knowledge_scope).value
        except ValueError as exc:
            raise HTTPException(400, str(exc)) from exc
    else:
        normalized_scope = ""
    store = await get_vector_store()
    results = []
    for filters in build_accessible_knowledge_filters():
        scoped_filters = {**filters, "source_file": doc_name}
        results.extend(await store.get_by_filter(scoped_filters, limit=10000))
    manageable = []
    for entry in results:
        metadata = dict(entry.metadata or {})
        visibility = str(metadata.get("visibility") or "private")
        if normalized_scope and visibility != normalized_scope:
            continue
        if metadata.get("source") != "user_upload":
            logger.warning("删除知识文档拒绝", doc_name=doc_name, reason="系统目录或内置文档")
            raise HTTPException(403, f"内置文档 '{doc_name}' 不可删除")
        resource_tenant_id = metadata.get("tenant_id")
        if resource_tenant_id is None and visibility != "system":
            resource_tenant_id = get_current_tenant_id()
        if can_manage_knowledge_resource(
            visibility,
            role=get_current_role(),
            current_tenant_id=get_current_tenant_id(),
            resource_tenant_id=resource_tenant_id,
            current_user_id=get_current_user_id(),
            owner_user_id=metadata.get("owner_user_id"),
        ):
            manageable.append(entry)
    if results and not manageable:
        logger.warning("删除知识文档权限拒绝", doc_name=doc_name, role=get_current_role())
        raise HTTPException(403, "无权删除该知识文档")
    # docs/metrics 仅包含内置文档，禁止通过管理 API 删除。
    base_dir = os.path.realpath(os.path.join(os.path.dirname(__file__), "..", "..", "docs", "metrics"))
    doc_path = os.path.realpath(os.path.join(base_dir, os.path.basename(doc_name)))
    if not doc_path.startswith(base_dir):
        raise HTTPException(403, "禁止访问")
    if os.path.isfile(doc_path):
        raise HTTPException(403, f"内置文档 '{doc_name}' 不可删除")
    # 从 PG 删除原始文件和 VectorStore 关联条目。
    from src.knowledge.file_store import get_file_store
    deleted_file = await get_file_store().delete(doc_name, knowledge_scope=normalized_scope)
    if manageable:
        await store.delete_by_ids([entry.id for entry in manageable])
    if not deleted_file and not manageable:
        raise HTTPException(404, f"文档 '{doc_name}' 未找到")
    result = {"status": "ok", "doc": doc_name}
    logger.info(
        "删除知识文档 API 完成",
        doc_name=doc_name,
        knowledge_scope=normalized_scope,
        vector_count=len(manageable),
        file_deleted=deleted_file,
    )
    return result


# ---- Skill 管理操作 ----


@router.post("/skills/refresh")
# 方法作用：重新扫描全部受信任 Skill 目录并返回当前身份可见数量。
# Args: 无。
# Returns: 刷新状态和可见 Skill 数量。
async def refresh_skills():
    """全量重新扫描 Skill 目录（用户手动触发）。"""
    try:
        from src.skill_manager import get_skill_manager
        await get_skill_manager().discover()
        from src.api.auth import get_current_tenant_id, get_current_user_id
        visible = get_skill_manager().get_visible_skills(
            get_current_tenant_id(), get_current_user_id(),
        )
        result = {"status": "ok", "total": len(visible)}
        logger.info("Skill 刷新完成", total=len(visible))
        return result
    except Exception as e:
        logger.error("Skill 刷新失败", error=str(e), exc_info=True)
        raise HTTPException(500, str(e))


@router.get("/skills/{skill_name}/content")
# 方法作用：读取当前身份可见 Skill 的原始 SKILL.md。
# Args: skill_name - Skill 名称；skill_scope - 可选精确作用域。
# Returns: Skill 名称、作用域和文件内容。
async def get_skill_content(skill_name: str, skill_scope: str | None = None):
    """获取 SKILL.md 文件的原始内容（直接从缓存 source_path 读取）。"""
    from src.skill_manager import get_skill_manager
    from src.api.auth import get_current_tenant_id, get_current_user_id
    mgr = get_skill_manager()
    skill = mgr.get_skill(
        skill_name, scope=skill_scope,
        tenant_id=get_current_tenant_id(), user_id=get_current_user_id(),
    )
    if not skill:
        raise HTTPException(404, f"Skill '{skill_name}' 未找到")
    md_path = os.path.join(skill.source_path, "SKILL.md")
    if not os.path.isfile(md_path):
        raise HTTPException(404, f"SKILL.md 文件不存在: {md_path}")
    with open(md_path, "r", encoding="utf-8", errors="replace") as f:
        result = {"name": skill_name, "scope": skill.scope, "content": f.read()}
    logger.info("读取 Skill 内容完成", name=skill_name, scope=skill.scope)
    return result


@router.put("/skills/{skill_name}/toggle")
# 方法作用：启用或停用当前身份有权管理的 Skill。
# Args: skill_name - Skill 名称；enabled - 目标状态；skill_scope - 可选精确作用域。
# Returns: 修改后的 Skill 状态。
async def toggle_skill(
    skill_name: str, enabled: bool = Query(...), skill_scope: str | None = None,
):
    """启用或禁用一个 Skill。"""
    try:
        from src.skill_manager import get_skill_manager
        from src.api.auth import (
            get_current_role, get_current_tenant_id, get_current_user_id,
        )
        from src.knowledge.governance import can_manage_knowledge_resource
        mgr = get_skill_manager()
        tenant_id = get_current_tenant_id()
        user_id = get_current_user_id()
        skill = mgr.get_skill(
            skill_name, scope=skill_scope, tenant_id=tenant_id, user_id=user_id,
        )
        if not skill:
            raise HTTPException(404, f"Skill '{skill_name}' 未找到")
        if not can_manage_knowledge_resource(
            skill.scope, role=get_current_role(), current_tenant_id=tenant_id,
            resource_tenant_id=skill.tenant_id, current_user_id=user_id,
            owner_user_id=skill.owner_user_id,
        ):
            raise HTTPException(403, "无权修改该 Skill")
        skill.enabled = enabled
        logger.info("Skill 启停完成", name=skill_name, scope=skill.scope, enabled=enabled)
        return {
            "status": "ok", "name": skill_name, "scope": skill.scope,
            "enabled": enabled,
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, str(e))


@router.delete("/skills/{skill_name}")
# 方法作用：删除当前身份有权管理的非内置 Skill 目录和缓存。
# Args: skill_name - Skill 名称；skill_scope - 可选精确作用域。
# Returns: 删除状态、名称和作用域。
async def delete_skill(skill_name: str, skill_scope: str | None = None):
    """删除一个 Skill（移除磁盘目录）。内置 Skill 不可删除。"""
    import shutil as _shutil
    from src.api.auth import get_current_role, get_current_tenant_id, get_current_user_id
    from src.knowledge.governance import can_manage_knowledge_resource
    from src.skill_manager import get_skill_manager
    mgr = get_skill_manager()
    tenant_id = get_current_tenant_id()
    user_id = get_current_user_id()
    skill = mgr.get_skill(
        skill_name, scope=skill_scope, tenant_id=tenant_id, user_id=user_id,
    )
    if not skill:
        raise HTTPException(404, f"Skill '{skill_name}' 目录不存在")
    if mgr.is_builtin(
        skill_name, tenant_id=tenant_id, user_id=user_id, scope=skill.scope,
    ):
        raise HTTPException(403, f"内置 Skill '{skill_name}' 不可删除")
    if not can_manage_knowledge_resource(
        skill.scope, role=get_current_role(), current_tenant_id=tenant_id,
        resource_tenant_id=skill.tenant_id, current_user_id=user_id,
        owner_user_id=skill.owner_user_id,
    ):
        raise HTTPException(403, "无权删除该 Skill")
    skill_dir = skill.source_path
    real = os.path.realpath(os.path.normpath(skill_dir))
    # 安全检查：必须在任一 allowed 目录下
    managed_root = os.path.realpath(os.path.normpath(str(mgr.managed_dir)))
    ok = os.path.commonpath([real, managed_root]) == managed_root
    if not ok:
        raise HTTPException(403, "路径非法")
    _shutil.rmtree(skill_dir)
    # 直接从缓存移除，不扫描整个文件夹
    get_skill_manager().remove_skill(
        skill_name, scope=skill.scope, tenant_id=tenant_id, user_id=user_id,
    )
    logger.info("Skill 删除完成", name=skill_name, scope=skill.scope)
    return {
        "status": "ok", "name": skill_name, "scope": skill.scope,
        "deleted": True,
    }


# ---- Health (11.1.10) ----

@router.get("/health", response_model=HealthResponse)
async def health():
    return HealthResponse(status="ok", llm_available=is_llm_available(),
                          uptime_seconds=round(time.time() - _started_at, 2))
